import traceback
import tkinter
import tkinter.messagebox
import sys
import os
from datetime import datetime

# --- CRASH REPORTER WRAPPER ---
try:
    import customtkinter as ctk
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from PIL import Image
    import openpyxl 

    # ==========================================
    # PART 1: DYNAMIC LOGIC
    # ==========================================

    def load_pricelist_and_packages():
        """
        Reads master_pricelist.xlsx.
        1. Builds the standard PRICELIST_DATA list for Data#3 logic.
        2. Builds the FITOUT_PACKAGES dictionary dynamically for Fit-Out logic.
        """
        script_dir = os.path.dirname(os.path.abspath(__file__))
        file_path = os.path.join(script_dir, "master_pricelist.xlsx")
        
        if not os.path.exists(file_path):
            return [], {}, "File 'master_pricelist.xlsx' not found."

        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            sheet = wb.active
            
            pricelist_data = []
            fitout_packages = {}

            # Iterate rows (skip header)
            for row in sheet.iter_rows(min_row=2, values_only=True):
                if row[0] is None and row[1] is None: continue

                # Map columns (0-based index)
                r_dist = float(row[0] or 0)
                r_name = str(row[1])
                r_vc_item = str(row[2]) if row[2] else ""
                r_disp_model = str(row[3])
                r_disp_price = float(row[4] or 0)
                r_mnt_model = str(row[5])
                r_mnt_price = float(row[6] or 0)
                r_cab_desc = str(row[7])
                r_cab_price = float(row[8] or 0)
                r_svc_price = float(row[9] or 0)
                r_ms_price = float(row[10] or 0)
                r_extras_str = str(row[11]) if len(row) > 11 and row[11] else ""

                # --- BUILD STANDARD DATA ENTRY (Data#3) ---
                cisco_list = [x.strip() for x in r_vc_item.split(',')]
                entry = {
                    "max_distance": r_dist,
                    "tier_name": r_name,
                    "cisco_items": cisco_list,
                    "display_model": r_disp_model,
                    "display_price": r_disp_price,
                    "mount_model": r_mnt_model,
                    "mount_price": r_mnt_price,
                    "cables_misc": r_cab_desc,
                    "cables_price": r_cab_price,
                    "service_price": r_svc_price,
                    "ms_annual": r_ms_price
                }
                pricelist_data.append(entry)

                # --- BUILD FIT-OUT PACKAGE (If applicable) ---
                if "Fit-Out" in r_name:
                    extra_items_list = []
                    if r_extras_str:
                        groups = r_extras_str.split(';')
                        for g in groups:
                            parts = g.split('|')
                            if len(parts) == 3:
                                extra_items_list.append(("Hardware", parts[0].strip(), float(parts[1]), int(parts[2])))

                    fitout_packages[r_name] = {
                        "display": ("Visual Display", r_disp_model, r_disp_price),
                        "mount": ("Mounting", r_mnt_model, r_mnt_price),
                        "vc": ("Video Conf", r_vc_item, 3900.00), 
                        "cables": ("Cabling", r_cab_desc, r_cab_price),
                        "services": r_svc_price,
                        "ms_price": r_ms_price,
                        "items": extra_items_list
                    }

            pricelist_data.sort(key=lambda x: x['max_distance'])
            return pricelist_data, fitout_packages, None

        except Exception as e:
            return [], {}, str(e)

    def get_room_configuration(distance, pricelist):
        if not pricelist: return None
        # IMPORTANT: When looking up Data#3 config, we must ignore Fit-Out rows in the pricelist
        for tier in pricelist:
            if "Fit-Out" in tier['tier_name']: 
                continue # Skip Fit-Out rows during Data#3 lookup
            if distance <= tier['max_distance']:
                return tier
        return None

    def generate_multi_room_proposal(client_name, room_list, project_mode, fitout_pkgs):
        
        # --- 1. SELECT TEMPLATE ---
        script_dir = os.path.dirname(os.path.abspath(__file__))
        
        if project_mode == "Data#3 (Cisco)":
            template_filename = "Template_Data3.docx"
        else:
            template_filename = "Template_Fitout.docx"

        template_path = os.path.join(script_dir, template_filename)
        
        if os.path.exists(template_path):
            doc = Document(template_path)
            doc.add_page_break() 
        else:
            doc = Document()

        # --- SET FONT TO HELVETICA ---
        try:
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Helvetica'
            font.size = Pt(10)
        except: pass 

        try:
            section = doc.sections[0]
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
        except: pass

        # --- HELPERS ---
        def shade_cell(cell, color_hex):
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
            cell._tc.get_or_add_tcPr().append(shading_elm)

        def add_manual_heading(text, size, color_rgb=None):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Helvetica'
            run.font.size = Pt(size)
            if color_rgb: run.font.color.rgb = color_rgb
            return p
            
        def format_row(row, height_cm):
            row.height = Cm(height_cm)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # ---------------------------------------------------------
        # PAGE 1: EXECUTIVE SUMMARY
        # ---------------------------------------------------------
        
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_t = p_title.add_run(f'AV Proposal: {client_name}')
        run_t.bold = True
        run_t.font.size = Pt(26)
        
        doc.add_paragraph(f'Prepared by: Alder Technology')
        doc.add_paragraph(f'Date: {datetime.now().strftime("%d/%m/%Y")}')
        doc.add_paragraph(f'Project Type: {project_mode}')
        doc.add_paragraph('------------------------------------------------------')

        add_manual_heading('Partnership Overview', 14)
        
        if project_mode == "Data#3 (Cisco)":
            overview_text = (
                "Alder Technology is pleased to partner with Data#3 to provide this solution. "
                "This document is split into two sections:\n"
                "1. A Master Financial Summary (Hardware + Year 1 Services).\n"
                "2. Detailed Bill of Materials for each specific room.\n\n"
                "Please note: Cisco hardware is listed for engineering reference but is to be supplied and priced by Data#3."
            )
        else:
            overview_text = (
                "Alder Technology is pleased to provide this comprehensive Audio Visual proposal for a complete office fit-out. "
                "This document outlines the Master Financial Summary and the Detailed Bill of Materials for every room, "
                "including all visual displays, conferencing bars, and installation services."
            )
            
        doc.add_paragraph(overview_text)

        add_manual_heading('1. Master Room Summary & Pricing', 14)
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        hdr_row = table.rows[0]
        format_row(hdr_row, 1.0)
        hdr = hdr_row.cells
        
        hdr[0].text = "Room Name"
        hdr[1].text = "Classification"
        hdr[2].text = "Supply & Services" 
        hdr[3].text = "Managed Service (Year 1)"
        hdr[4].text = "Total Year 1 (Ex GST)"

        for cell in hdr:
            shade_cell(cell, "D9E2F3")
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        grand_total_project = 0

        # --- CALCULATE TOTALS ---
        for room in room_list:
            name = room['name']
            r_type = room['type'] # "Fit-Out 55" or "Medium"
            dist = room['distance'] 

            # --- FIT OUT CALCULATION ---
            if "Fit-Out" in r_type:
                pkg = fitout_pkgs[r_type]
                
                ms_annual = pkg['ms_price']
                
                # Handle Dual Screen Cost
                qty_display = 2 if "Dual" in r_type else 1
                cost_display = pkg['display'][2] * qty_display
                
                cost_mount = pkg['mount'][2]
                cost_vc = pkg['vc'][2]
                cost_cables = pkg['cables'][2]
                cost_services = pkg['services']
                
                cost_extras = sum([item[2] * item[3] for item in pkg['items']])

                upfront_cost = cost_display + cost_mount + cost_vc + cost_cables + cost_services + cost_extras
                display_label = f"{r_type} ({dist}m)"
                
            # --- DATA#3 CALCULATION ---
            else:
                data = room['config']
                upfront_cost = data['display_price'] + data['mount_price'] + data['cables_price'] + data['service_price']
                ms_annual = data['ms_annual']
                display_label = f"{data['tier_name']} ({dist}m)"

            room_total_y1 = upfront_cost + ms_annual
            grand_total_project += room_total_y1

            row_obj = table.add_row()
            format_row(row_obj, 0.9)
            row = row_obj.cells
            
            row[0].text = name
            row[1].text = display_label
            row[2].text = f"${upfront_cost:,.0f}"
            row[3].text = f"${ms_annual:,.0f}"
            row[4].text = f"${room_total_y1:,.2f}"
            
            for i in range(2, 5):
                row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph("\n")
        p_total = doc.add_paragraph()
        p_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        runner = p_total.add_run(f"TOTAL YEAR 1 PROJECT VALUE (EX GST): ${grand_total_project:,.2f}")
        runner.bold = True
        runner.font.size = Pt(16)
        runner.font.color.rgb = RGBColor(0, 102, 204)
        
        # ---------------------------------------------------------
        # PAGE 2+: DETAILED ROOM BREAKDOWNS
        # ---------------------------------------------------------
        
        doc.add_page_break()
        add_manual_heading('2. Detailed Room Specifications', 16)

        for room in room_list:
            name = room['name']
            r_type = room['type']
            dist = room['distance']
            
            table_room = doc.add_table(rows=0, cols=5)
            table_room.style = 'Table Grid'
            table_room.autofit = False 
            table_room.columns[0].width = Cm(3.0)
            table_room.columns[1].width = Cm(1.3)
            table_room.columns[2].width = Cm(8.5)
            table_room.columns[3].width = Cm(2.5)
            table_room.columns[4].width = Cm(2.7)
            
            # --- HEADER ---
            row_hdr_obj = table_room.add_row()
            format_row(row_hdr_obj, 1.0)
            row_hdr = row_hdr_obj.cells
            row_hdr[0].merge(row_hdr[4])
            row_hdr[0].text = f"ROOM: {name} - {r_type} (Furthest Participant: {dist}m)"
            shade_cell(row_hdr[0], "1F4E79")
            row_hdr[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            row_hdr[0].paragraphs[0].runs[0].bold = True

            # --- IMAGE PLACEHOLDER ---
            row_img_obj = table_room.add_row()
            format_row(row_img_obj, 4.0)
            row_img = row_img_obj.cells
            row_img[0].merge(row_img[4])
            row_img[0].text = "[PASTE FLOOR PLAN IMAGE HERE]"
            row_img[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # --- COL HEADERS ---
            row_cols_obj = table_room.add_row()
            format_row(row_cols_obj, 0.9)
            row_cols = row_cols_obj.cells
            row_cols[0].text = "Item"
            row_cols[1].text = "Qty"
            row_cols[2].text = "Description / Model"
            row_cols[3].text = "Unit Cost"
            row_cols[4].text = "Total"
            for c in row_cols:
                shade_cell(c, "D9E2F3")
                c.paragraphs[0].runs[0].bold = True

            # =========================================================
            # MODE A: FIT-OUT LOGIC
            # =========================================================
            if "Fit-Out" in r_type:
                pkg = fitout_pkgs[r_type]
                
                row_sec1_obj = table_room.add_row()
                format_row(row_sec1_obj, 0.8)
                row_sec1 = row_sec1_obj.cells
                row_sec1[0].merge(row_sec1[4])
                row_sec1[0].text = "1. Hardware & Services Scope"
                shade_cell(row_sec1[0], "E7E6E6")
                row_sec1[0].paragraphs[0].runs[0].bold = True

                def add_row(cat, qty, desc, price):
                    r_obj = table_room.add_row()
                    format_row(r_obj, 0.9)
                    r = r_obj.cells
                    r[0].text = cat
                    r[1].text = str(qty)
                    r[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r[2].text = desc
                    r[3].text = f"${price:,.2f}"
                    r[4].text = f"${(price * qty):,.2f}"
                    r[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    r[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

                # 1. VC BAR
                add_row(pkg['vc'][0], 1, pkg['vc'][1], pkg['vc'][2])
                
                # 2. DISPLAY
                disp_qty = 2 if "Dual" in r_type else 1
                add_row(pkg['display'][0], disp_qty, pkg['display'][1], pkg['display'][2])
                
                # 3. EXTRA ITEMS
                for item in pkg['items']:
                    add_row(item[0], item[3], item[1], item[2])
                
                # 4. MOUNT & CABLES
                add_row(pkg['mount'][0], 1, pkg['mount'][1], pkg['mount'][2])
                add_row(pkg['cables'][0], 1, pkg['cables'][1], pkg['cables'][2])

                # 6. SERVICES (FIXED)
                r_svc_obj = table_room.add_row()
                format_row(r_svc_obj, 0.9)
                r_svc = r_svc_obj.cells
                r_svc[0].text = "Services"
                r_svc[1].text = "1"
                r_svc[2].text = "Total Services (Staging, Installation, PM, Engineering)"
                r_svc[3].text = f"${pkg['services']:,.2f}"
                r_svc[4].text = f"${pkg['services']:,.2f}"
                r_svc[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r_svc[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

                # --- MANAGED SERVICES ---
                row_sec3_obj = table_room.add_row()
                format_row(row_sec3_obj, 0.8)
                row_sec3 = row_sec3_obj.cells
                row_sec3[0].merge(row_sec3[4])
                row_sec3[0].text = "2. Managed Services"
                shade_cell(row_sec3[0], "E7E6E6")
                row_sec3[0].paragraphs[0].runs[0].bold = True

                ms_val = pkg['ms_price']
                
                r_msa_obj = table_room.add_row()
                format_row(r_msa_obj, 0.9)
                r_msa = r_msa_obj.cells
                r_msa[0].text = "Support"
                r_msa[1].text = "1"
                r_msa[2].text = "Managed Service Agreement - Year 1 (Annual Billing)"
                r_msa[3].text = f"${ms_val:,.2f}"
                r_msa[4].text = f"${ms_val:,.2f}"
                r_msa[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r_msa[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

                # CALC TOTAL
                disp_total = pkg['display'][2] * disp_qty
                extra_total = sum([item[2] * item[3] for item in pkg['items']])
                total_room_cost = disp_total + pkg['mount'][2] + pkg['vc'][2] + pkg['cables'][2] + pkg['services'] + extra_total + ms_val

            # =========================================================
            # MODE B: DATA#3 (CISCO) LOGIC
            # =========================================================
            else:
                data = room['config']
                
                final_cisco_list = []
                moved_to_alder_list = []

                for item in data['cisco_items']:
                    if not item: continue
                    clean_item = item.strip()
                    if "Shure" in clean_item:
                        moved_to_alder_list.append(clean_item)
                    else:
                        item_desc = clean_item
                        if "Room Bar Pro" in clean_item: item_desc += " / CS-BARPRO-K9"
                        elif "Room Bar" in clean_item and "Pro" not in clean_item: item_desc += " / CS-BAR-T-C-K9"
                        elif "Ceiling Microphone Pro" in clean_item: item_desc += " / CS-MIC-CLGPRO="
                        elif "Wire Hanging" in clean_item: item_desc += " / CS-MIC-CLGP-WHK="
                        final_cisco_list.append(item_desc)

                if final_cisco_list:
                    row_sec1_obj = table_room.add_row()
                    format_row(row_sec1_obj, 0.8)
                    row_sec1 = row_sec1_obj.cells
                    row_sec1[0].merge(row_sec1[4])
                    row_sec1[0].text = "1. Data#3 Supply Scope (Cisco Hardware)"
                    shade_cell(row_sec1[0], "E7E6E6")
                    row_sec1[0].paragraphs[0].runs[0].bold = True

                    for item in final_cisco_list:
                        r_obj = table_room.add_row()
                        format_row(r_obj, 0.9)
                        r = r_obj.cells
                        for cell in r: shade_cell(cell, "FFF2CC")
                        r[0].text = "Video Conf"
                        r[1].text = "1"
                        r[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r[2].text = item 
                        r[3].text = "Excl."
                        r[4].text = "Excl."
                        r[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        r[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

                row_sec2_obj = table_room.add_row()
                format_row(row_sec2_obj, 0.8)
                row_sec2 = row_sec2_obj.cells
                row_sec2[0].merge(row_sec2[4])
                row_sec2[0].text = "2. Alder Technology Supply Scope"
                shade_cell(row_sec2[0], "E7E6E6")
                row_sec2[0].paragraphs[0].runs[0].bold = True

                def add_spec_row(item_cat, qty, desc_text, price_val):
                    r_obj = table_room.add_row()
                    format_row(r_obj, 0.9)
                    r = r_obj.cells
                    r[0].text = item_cat
                    r[1].text = str(qty)
                    r[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r[2].text = desc_text
                    r[3].text = f"${price_val:,.2f}"
                    r[4].text = f"${(price_val * qty):,.2f}"
                    r[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    r[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

                add_spec_row("Visual Display", 1, data['display_model'], data['display_price'])
                add_spec_row("Mounting", 1, data['mount_model'], data['mount_price'])
                for moved_item in moved_to_alder_list:
                    add_spec_row("Conf/Audio", 1, moved_item, 0.00) 
                add_spec_row("Cabling", 1, data['cables_misc'], data['cables_price'])
                
                r_svc_obj = table_room.add_row()
                format_row(r_svc_obj, 0.9)
                r_svc = r_svc_obj.cells
                r_svc[0].text = "Services"
                r_svc[1].text = "1"
                r_svc[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_svc[2].text = "Professional Services: Installation, Staging & PM"
                r_svc[3].text = f"${data['service_price']:,.2f}"
                r_svc[4].text = f"${data['service_price']:,.2f}"
                r_svc[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r_svc[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

                row_sec3_obj = table_room.add_row()
                format_row(row_sec3_obj, 0.8)
                row_sec3 = row_sec3_obj.cells
                row_sec3[0].merge(row_sec3[4])
                row_sec3[0].text = "3. Managed Services"
                shade_cell(row_sec3[0], "E7E6E6")
                row_sec3[0].paragraphs[0].runs[0].bold = True

                r_msa_obj = table_room.add_row()
                format_row(r_msa_obj, 0.9)
                r_msa = r_msa_obj.cells
                r_msa[0].text = "Support"
                r_msa[1].text = "1"
                r_msa[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_msa[2].text = "Managed Service Agreement - Year 1 (Annual Billing)"
                r_msa[3].text = f"${data['ms_annual']:,.2f}"
                r_msa[4].text = f"${data['ms_annual']:,.2f}"
                r_msa[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r_msa[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

                upfront_subtotal = data['display_price'] + data['mount_price'] + data['cables_price'] + data['service_price']
                total_room_cost = upfront_subtotal + data['ms_annual']

            # --- TOTAL ROW ---
            row_sub_obj = table_room.add_row()
            format_row(row_sub_obj, 1.0)
            row_sub = row_sub_obj.cells
            row_sub[0].merge(row_sub[3])
            row_sub[0].text = "TOTAL (EX GST):"
            row_sub[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_sub[0].paragraphs[0].runs[0].bold = True
            row_sub[4].text = f"${total_room_cost:,.2f}"
            row_sub[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_sub[4].paragraphs[0].runs[0].bold = True

            doc.add_paragraph("")

        # ---------------------------------------------------------
        # FINAL SECTION
        # ---------------------------------------------------------
        doc.add_page_break()
        add_manual_heading('Managed Service Agreement', 14)
        doc.add_paragraph(
            "Pricing excludes GST and is charged annually with an increase each year of 4% or CPI whichever is the greater. "
            "Acceptance of a 60 month agreement upfront locks pricing for the five year term with no increase for CPI. "
            "Pricing includes all cloud monitoring hosting charges and any onsite support required."
        )

        add_manual_heading('Exclusions', 14)
        exclusions_text = (
            "We exclude the following items from our installation. "
            "We exclude all power and data, plus building works. All works required shall be identified and must be completed prior to our installers attending site. "
            "Any Cat6/6A Tie Lines between Audio Visual Locations (e.g. behind display to table box/floor box) shall be by the data contractor. "
            "We exclude all decommissioning, disposal of equipment and make good works. "
            "We exclude all height access equipment, and all furniture protection equipment/coverings. "
            "All Webex/Teams/Exchange credentials must be provided prior to attending site. Admin credentials for any Webex/Teams endpoint must be provided to Alder Technology for the duration of any deployment including temporary administrator cloud tenant access and other software admin access for the duration of deployment. "
            "The network must be active and configured for Teams/Webex prior to attending site. "
            "All external services provided by the client or their nominated systems integrator, including Microsoft Teams Room/Webex accounts, Azure Intune registrations and specific deployment requirements, Exchange, Skype for Business and Microsoft security and compliance requirements and Fast track or Peering ISP plans are the responsibility of the client. "
            "Software updates and the impact on the hardware, user experience and the operational state of the system are the sole responsibility of the client. Any such updates that require Alder Technology site attendance incur a Service call out fee and hourly rate at agreed hourly rates, unless a service level agreement covering these works is in place. "
            "A site planner shall be provided by Alder Technology and must be completed by the client prior to our installers attending site. "
            "Should works not be able to commence due to the above or client led delays, a call out fee may be applied. "
            "Pricing is based on a consecutive work package. Significant delays or pauses in works due to client delays or room unavailability may result in additional charges."
        )
        doc.add_paragraph(exclusions_text)

        doc.add_paragraph("\nThank you for your consideration. Please call me if you have any further queries.")
        doc.add_paragraph("Regards,")
        sig = doc.add_paragraph("George Coles")
        sig.runs[0].bold = True

        desktop = os.path.expanduser("~/Desktop")
        save_folder = os.path.join(desktop, "Alder_Quotes")
        if not os.path.exists(save_folder):
            os.makedirs(save_folder)

        safe_client_name = "".join([c for c in client_name if c.isalpha() or c.isdigit() or c==' ']).rstrip()
        timestamp = datetime.now().strftime("%H-%M-%S")
        filename = f"Alder_Quote_{safe_client_name}_{project_mode[:4]}_{timestamp}.docx"
        full_path = os.path.join(save_folder, filename)
        
        doc.save(full_path)
        return full_path

    # ==========================================
    # PART 2: THE USER INTERFACE (GUI)
    # ==========================================

    ctk.set_appearance_mode("Light") 
    ctk.set_default_color_theme("green") 

    app = ctk.CTk()
    app.geometry("600x900") 
    app.title("Alder Technology - Quoting Tool")
    app.configure(fg_color="#F0F0F0") # Cleaner background

    PRICELIST_DATA, FITOUT_PACKAGES_DYN, EXCEL_ERROR = load_pricelist_and_packages()
    ADDED_ROOMS = [] 

    # --- LOGO ---
    script_dir = os.path.dirname(os.path.abspath(__file__))
    logo_path = os.path.join(script_dir, "alder_logo.png")
    logo_image = None
    if os.path.exists(logo_path):
        try:
            pil_image = Image.open(logo_path)
            logo_image = ctk.CTkImage(light_image=pil_image, dark_image=pil_image, size=(280, 200))
        except Exception: pass

    # --- REFRESH ROOM LIST ---
    def refresh_room_list():
        for widget in scroll_frame.winfo_children():
            widget.destroy()

        if not ADDED_ROOMS:
            lbl_empty = ctk.CTkLabel(scroll_frame, text="No rooms added yet.", text_color="gray")
            lbl_empty.pack(pady=20)
            return

        for index, room in enumerate(ADDED_ROOMS):
            row_frame = ctk.CTkFrame(scroll_frame, fg_color="white", corner_radius=5)
            row_frame.pack(fill="x", pady=2, padx=5)

            if "Fit-Out" in room['type']:
                display_txt = f"📦 {room['name']} ({room['type']} - FP: {room['distance']}m)"
            else:
                display_txt = f"🔌 {room['name']} ({room['type']} - FP: {room['distance']}m)"

            lbl_info = ctk.CTkLabel(row_frame, text=display_txt, anchor="w", font=("Arial", 12))
            lbl_info.pack(side="left", padx=10, pady=5)

            btn_del = ctk.CTkButton(
                row_frame, text="✖", width=30, height=25, 
                fg_color="#FFEEEE", text_color="red", hover_color="#FFDDDD",
                command=lambda i=index: delete_room(i)
            )
            btn_del.pack(side="right", padx=5)

    def delete_room(index):
        if 0 <= index < len(ADDED_ROOMS):
            del ADDED_ROOMS[index]
            refresh_room_list()

    def update_dropdown_options(choice):
        """ Updates the room type dropdown based on project mode """
        if choice == "Data#3 (Cisco)":
            # FILTER: Exclude rows with 'Fit-Out' in name
            options = [item['tier_name'] for item in PRICELIST_DATA if "Fit-Out" not in item['tier_name']]
        else:
            # FILTER: Show ONLY Fit-Out packages
            options = list(FITOUT_PACKAGES_DYN.keys()) if FITOUT_PACKAGES_DYN else ["No Fit-Out Pkgs in Excel"]
            
        entry_distance.configure(state="normal", placeholder_text="Furthest Participant (m)")
        entry_distance.delete(0, "end")
            
        dropdown_type.configure(values=options)
        dropdown_type.set(options[0] if options else "No Options")

    def on_add_room():
        r_name = entry_room_name.get().strip()
        r_type = dropdown_type.get()
        mode = dropdown_project_mode.get()

        if not r_name:
            tkinter.messagebox.showwarning("Missing Data", "Please enter a Room Name.")
            return

        try:
            r_dist_str = entry_distance.get().strip()
            if not r_dist_str:
                tkinter.messagebox.showwarning("Missing Data", "Furthest Participant distance is required.")
                return
            r_dist = float(r_dist_str)
        except ValueError:
            tkinter.messagebox.showwarning("Invalid Distance", "Please enter a valid number for distance.")
            return

        room_entry = {'name': r_name, 'type': r_type, 'distance': r_dist}

        if mode == "Data#3 (Cisco)":
            config = get_room_configuration(r_dist, PRICELIST_DATA)
            if not config:
                tkinter.messagebox.showerror("Error", f"Distance {r_dist}m exceeds Excel tiers for Data#3.")
                return
            room_entry['config'] = config
        else:
            room_entry['config'] = None
        
        ADDED_ROOMS.append(room_entry)
        entry_room_name.delete(0, "end")
        refresh_room_list()

    def on_generate_click():
        lbl_status.configure(text="Processing...", text_color="black")
        app.update()
        
        if EXCEL_ERROR:
            lbl_status.configure(text=f"EXCEL ERROR: {EXCEL_ERROR}", text_color="#FF5555")
            return

        client = entry_client.get().strip()
        if not client:
            lbl_status.configure(text="Error: Enter Client Name", text_color="#FF5555")
            return

        if len(ADDED_ROOMS) == 0:
            lbl_status.configure(text="Error: Please add at least one room.", text_color="#FF5555")
            return

        selected_mode = dropdown_project_mode.get()

        try:
            filepath = generate_multi_room_proposal(client, ADDED_ROOMS, selected_mode, FITOUT_PACKAGES_DYN)
            lbl_status.configure(text=f"SUCCESS!\nSaved to Desktop/Alder_Quotes:\n{os.path.basename(filepath)}", text_color="#009A44")
            try: os.startfile(os.path.dirname(filepath))
            except: pass
        except PermissionError:
            lbl_status.configure(text="ERROR: PERMISSION DENIED.\nClose Word and try again.", text_color="#FF5555")
        except Exception as e:
            lbl_status.configure(text=f"Error: {str(e)}", text_color="#FF5555")

    # --- GUI LAYOUT ---
    
    # Main container
    main_frame = ctk.CTkFrame(app, corner_radius=0, fg_color="white")
    main_frame.pack(fill="both", expand=True)

    # Header Section
    header_frame = ctk.CTkFrame(main_frame, fg_color="#F9F9F9")
    header_frame.pack(fill="x", pady=0)
    
    if logo_image:
        lbl_logo = ctk.CTkLabel(header_frame, image=logo_image, text="")
        lbl_logo.pack(pady=15)
    else:
        lbl_title = ctk.CTkLabel(header_frame, text="ALDER TECHNOLOGY", font=("Arial Black", 24), text_color="#009A44")
        lbl_title.pack(pady=20)

    # --- EXCEL STATUS ---
    if EXCEL_ERROR:
        lbl_db_status = ctk.CTkLabel(main_frame, text=f"⚠️ DB ERROR: {EXCEL_ERROR}", text_color="red")
    else:
        lbl_db_status = ctk.CTkLabel(main_frame, text=f"✔ Database Connected ({len(PRICELIST_DATA)} tiers)", text_color="green", font=("Arial", 10))
    lbl_db_status.pack(pady=(5, 10))

    # --- INPUT SECTION ---
    input_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
    input_frame.pack(pady=10)

    # Client Name
    ctk.CTkLabel(input_frame, text="Client Name:", font=("Arial", 12, "bold")).grid(row=0, column=0, sticky="w", padx=10)
    entry_client = ctk.CTkEntry(input_frame, placeholder_text="e.g. Acme Corp", width=300)
    entry_client.grid(row=0, column=1, pady=5)

    # Project Mode
    ctk.CTkLabel(input_frame, text="Project Mode:", font=("Arial", 12, "bold")).grid(row=1, column=0, sticky="w", padx=10)
    dropdown_project_mode = ctk.CTkOptionMenu(
        input_frame,
        values=["Data#3 (Cisco)", "Fit-Out (Full Scope)"],
        width=300,
        command=update_dropdown_options,
        fg_color="#333", button_color="#222"
    )
    dropdown_project_mode.set("Data#3 (Cisco)")
    dropdown_project_mode.grid(row=1, column=1, pady=5)

    # Separator
    ctk.CTkFrame(main_frame, height=2, fg_color="#EEE").pack(fill="x", pady=15, padx=20)

    # --- ROOM BUILDER SECTION ---
    builder_frame = ctk.CTkFrame(main_frame, fg_color="#F4F4F4", corner_radius=10)
    builder_frame.pack(pady=0, padx=20, fill="x")

    ctk.CTkLabel(builder_frame, text="Room Builder", font=("Arial", 14, "bold"), text_color="#555").pack(pady=(10, 5))

    # Room Type Dropdown
    dropdown_type = ctk.CTkOptionMenu(builder_frame, values=[], width=350, fg_color="#1F4E79", button_color="#143656")
    dropdown_type.pack(pady=5)
    
    # Room Name
    entry_room_name = ctk.CTkEntry(builder_frame, placeholder_text="Room Name (e.g. Boardroom)", width=350)
    entry_room_name.pack(pady=5)

    # Furthest Participant
    entry_distance = ctk.CTkEntry(builder_frame, placeholder_text="Furthest Participant (m)", width=350)
    entry_distance.pack(pady=5)

    btn_add = ctk.CTkButton(builder_frame, text="+ ADD ROOM", width=350, fg_color="#666", hover_color="#444", command=on_add_room)
    btn_add.pack(pady=(10, 15))

    # --- LIST SECTION ---
    list_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
    list_frame.pack(pady=10, padx=20, fill="both", expand=True)

    ctk.CTkLabel(list_frame, text="Added Rooms:", font=("Arial", 12, "bold")).pack(anchor="w")
    scroll_frame = ctk.CTkScrollableFrame(list_frame, height=200, corner_radius=10, fg_color="#EEE")
    scroll_frame.pack(fill="both", expand=True, pady=5)

    # --- GENERATE BTN ---
    btn_generate = ctk.CTkButton(
        main_frame, text="GENERATE PROPOSAL", command=on_generate_click, 
        width=300, height=50, font=("Arial", 15, "bold"),
        fg_color="#009A44", hover_color="#007a36"
    )
    btn_generate.pack(pady=20)

    lbl_status = ctk.CTkLabel(main_frame, text="Ready", text_color="black")
    lbl_status.pack(pady=(0, 20))
    
    # Initialize Dropdowns
    update_dropdown_options("Data#3 (Cisco)")
    refresh_room_list()

    app.mainloop()

except Exception as e:
    root = tkinter.Tk()
    root.withdraw()
    error_msg = str(traceback.format_exc())
    tkinter.messagebox.showerror("Program Error", f"The program crashed:\n\n{error_msg}")
