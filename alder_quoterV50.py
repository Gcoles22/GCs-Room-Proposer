import traceback
import tkinter
import tkinter.messagebox
import sys
import os
from datetime import datetime

# --- CRASH REPORTER WRAPPER ---
try:
    import customtkinter as ctk
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from PIL import Image

    # ==========================================
    # PART 1: HARD-CODED DATA & LOGIC
    # ==========================================

    def load_internal_data():
        """
        REPLACES EXCEL LOADING.
        Returns hard-coded dictionaries for Data#3 and Fit-Out packages.
        """
        
        # ---------------------------------------------------------
        # 1. FIT-OUT PACKAGES
        # ---------------------------------------------------------
        
        # 98" AUDIO UPGRADE DATA
        audio_upgrade_98 = {
            "name": "Premium Audio Upgrade (Xilica/Sennheiser)",
            "total_price": 15860.00, 
            "items": [
                ("DSP", 1, "Xilica Room Hub - AI Based Digital Signal Processor. Dual NIC, USB, 8ch AEC", 3200.00),
                ("Mic", 1, "Sennheiser TeamConnect Ceiling 2 - Beamforming Microphone POE, White", 5800.00),
                ("Speakers", 4, "Xilica Sonia-C5 - Bezel-less 5.25” coaxial in-ceiling loudspeaker", 240.00),
                ("Amp", 1, "Xilica Sonia-Amp - Four Channel POE++ Amplifier with Dante", 1300.00),
                ("Cabling", 1, "Custom Cables, Hardware and Consumables", 600.00),
                ("Services", 1, "Project Services (Staging, Install, Engineering, PM)", 4000.00)
            ]
        }

        # UPDATED MS PRICING: 0-6.5m = $1200, 7.5m = $1500
        fitout_packages = {
            "Fit-Out 55": {
                "max_distance": 3.0,
                "display": ("Visual Display", "LG 55UL3J-B - 55\" UHD Commercial Display", 1600.00),
                "mount": ("Mounting", "Wall Mount Bracket (Tilt)", 80.00),
                "vc": ("Video Conf", "Maxhub XBAR W70", 3900.00), 
                "cables": ("Cabling", "Cables, Hardware and Consumables", 200.00),
                "services": 2500.00,
                "ms_price": 1200.00, # Updated
                "items": []
            },
            "Fit-Out 65": {
                "max_distance": 4.5,
                "display": ("Visual Display", "LG 65UL3J-B - 65\" UHD Commercial Display", 2100.00),
                "mount": ("Mounting", "Wall Mount Bracket (Tilt)", 90.00),
                "vc": ("Video Conf", "Maxhub XBAR W70", 3900.00), 
                "cables": ("Cabling", "Cables, Hardware and Consumables", 250.00),
                "services": 2800.00,
                "ms_price": 1200.00, # Updated
                "items": []
            },
            "Fit-Out 75": {
                "max_distance": 5.5,
                "display": ("Visual Display", "LG 75UL3J-B - 75\" UHD Commercial Display", 2800.00),
                "mount": ("Mounting", "Wall Mount Bracket Heavy Duty", 100.00),
                "vc": ("Video Conf", "Maxhub XBAR W70", 3900.00), 
                "cables": ("Cabling", "Cables, Hardware and Consumables", 300.00),
                "services": 3000.00,
                "ms_price": 1200.00, # Updated
                "items": []
            },
            "Fit-Out 86": {
                "max_distance": 6.5,
                "display": ("Visual Display", "LG 86UL3J-B - Commercial Professional Monitor 86\" LED, 4K UHD", 3300.00),
                "mount": ("Mounting", "Wall Mount Bracket VP-F100 (82\"-98\")", 100.00),
                "vc": ("Video Conf", "Maxhub XBAR W70 - Teams Certified Windows 11 MTR", 3900.00), 
                "cables": ("Cabling", "Custom Cables, Hardware and Consumables", 300.00),
                "services": 3500.00, 
                "ms_price": 1200.00, # Updated
                "items": []
            },
            "Fit-Out 98": {
                "max_distance": 7.5,
                "display": ("Visual Display", "LG 98UM5K - Commercial Professional Monitor 98\" LED, 4K UHD", 7500.00),
                "mount": ("Mounting", "Wall Mount Bracket VP-F100 (82\"-98\")", 100.00),
                "vc": ("Video Conf", "Maxhub XBAR W70 - Teams Certified Windows 11 MTR", 3900.00), 
                "cables": ("Cabling", "Custom Cables, Hardware and Consumables", 300.00),
                "services": 3500.00,
                "ms_price": 1500.00, # Updated
                "items": [],
                "audio_upgrade": audio_upgrade_98
            }
        }

        # ---------------------------------------------------------
        # 2. DATA#3 PACKAGES
        # ---------------------------------------------------------
        pricelist_data = [
            {
                "max_distance": 3.0,
                "tier_name": "Small Room",
                "cisco_items": ["Cisco Room Bar"],
                "display_model": "Samsung 55\" Commercial Display",
                "display_price": 1600.00,
                "mount_model": "Wall Mount Bracket",
                "mount_price": 100.00,
                "cables_misc": "HDMI & Patch Leads",
                "cables_price": 150.00,
                "service_price": 1500.00,
                "ms_annual": 1200.00
            },
            {
                "max_distance": 4.5,
                "tier_name": "Medium Room",
                "cisco_items": ["Cisco Room Bar", "Cisco Table Microphone Pro"],
                "display_model": "Samsung 65\" Commercial Display",
                "display_price": 2100.00,
                "mount_model": "Wall Mount Bracket",
                "mount_price": 120.00,
                "cables_misc": "HDMI & Patch Leads",
                "cables_price": 200.00,
                "service_price": 1800.00,
                "ms_annual": 1200.00
            },
            {
                "max_distance": 5.5,
                "tier_name": "Large Room",
                "cisco_items": ["Cisco Room Bar Pro", "Cisco Ceiling Microphone Pro"],
                "display_model": "Samsung 75\" Commercial Display",
                "display_price": 2800.00,
                "mount_model": "Heavy Duty Wall Mount",
                "mount_price": 180.00,
                "cables_misc": "Integration Kit & Cabling",
                "cables_price": 250.00,
                "service_price": 2500.00,
                "ms_annual": 1200.00
            },
            {
                "max_distance": 6.5,
                "tier_name": "X-Large Room",
                "cisco_items": ["Cisco Room Bar Pro", "Cisco Ceiling Microphone Pro"],
                "display_model": "Samsung 85\"/86\" Commercial Display",
                "display_price": 3500.00,
                "mount_model": "Heavy Duty Wall Mount (86\")",
                "mount_price": 250.00,
                "cables_misc": "Integration Kit & Cabling",
                "cables_price": 300.00,
                "service_price": 2800.00,
                "ms_annual": 1200.00
            },
            {
                "max_distance": 7.5,
                "tier_name": "Boardroom",
                "cisco_items": [
                    "Cisco Room Kit EQ", 
                    "Cisco Quad Camera", 
                    "2x Cisco Ceiling Microphone Pro", 
                    "AV Integrator License",
                    "6-8x Shure Ceiling Speakers"
                ],
                "display_model": "Samsung 98\" Commercial Display",
                "display_price": 9500.00,
                "mount_model": "Heavy Duty Wall Mount (98\")",
                "mount_price": 350.00,
                "cables_misc": "Audio Integration & Cabling Kit",
                "cables_price": 600.00,
                "service_price": 4500.00, 
                "ms_annual": 1500.00
            }
        ]

        # Sort for safety
        pricelist_data.sort(key=lambda x: x['max_distance'])
        
        return pricelist_data, fitout_packages, None

    def get_fitout_text_blocks(r_type):
        """
        Returns a list of tuples: (Heading, BodyText)
        """
        # --- 55" ROOM (Up to 3.0m) ---
        if "55" in r_type:
            return [
                ("Proposed Solution",
                 "The 6P meeting room represents rooms with a maximum viewing distance of up to 3m. Each 6P meeting room shall use a 55” display, with a Maxhub W70 Bar. This Windows 11 based unit has a quad camera, best in class AI based audio and robust cloud based monitoring.\nAt the table, USB-C connectivity is included for content sharing. BYOD is also possible via this cable. It is recommended that MS Teams shall be the primary connectivity method using the inbuilt wireless connectivity."),
                ("Works in Association",
                 "Behind the LCD, mounted offset to avoid the LCD bracket, there will need to be 1x Double GPO, and two data points (Teams Compute, Display). A Cat6A cable will need to be run from behind the display to the table box for the touch screen console and content sharing. Should BYOD be required, a second Cat6A should be run from behind the display to the table box."),
                ("Room Options",
                 "The room has the option of a room booking panel. This is a Teams certified room booking panel, that allows for users to see the status of the room (red for occupied, green for available), book the room from the touch screen, and also book it as a Teams meeting. Room booking panels may have light bars added for easy identification of room availability, and occupancy sensors, which release unoccupied rooms from unused bookings.")
            ]

        # --- 65" ROOM (Up to 4.5m) ---
        elif "65" in r_type:
            return [
                ("Proposed Solution", 
                 "The 6P meeting room represents rooms with a maximum viewing distance of up to 4.5m. Each 6P meeting room shall use a 65” display, with a Maxhub W70 Bar. This Windows 11 based unit has a quad camera, best in class AI based audio and robust cloud based monitoring.\nAt the table, USB-C connectivity is included for content sharing. BYOD is also possible via this cable. It is recommended that MS Teams shall be the primary connectivity method using the inbuilt wireless connectivity."),
                ("Works in Association", 
                 "Behind the LCD, mounted offset to avoid the LCD bracket, there will need to be 1x Double GPO, and two data points (Teams Compute, Display). A Cat6A cable will need to be run from behind the display to the table box for the touch screen console and content sharing. Should BYOD be required, a second Cat6A should be run from behind the display to the table box."),
                ("Room Options", 
                 "The room has the option of a room booking panel. This is a Teams certified room booking panel, that allows for users to see the status of the room (red for occupied, green for available), book the room from the touch screen, and also book it as a Teams meeting. Room booking panels may have light bars added for easy identification of room availability, and occupancy sensors, which release unoccupied rooms from unused bookings.")
            ]

        # --- 75" ROOM (Up to 5.5m) ---
        elif "75" in r_type:
            return [
                ("Proposed Solution", 
                 "The 8P meeting room represents rooms with a maximum viewing distance of up to 5.5m. Each 8P meeting room shall use a 75” display, with a Maxhub W70 Bar. This Windows 11 based unit has a quad camera, best in class AI based audio and robust cloud based monitoring.\nAt the table, USB-C connectivity is included for content sharing. BYOD is also possible via this cable. It is recommended that MS Teams shall be the primary connectivity method using the inbuilt wireless connectivity."),
                ("Works in Association", 
                 "Behind the LCD, mounted offset to avoid the LCD bracket, there will need to be 1x Double GPO, and two data points (Teams Compute, Display). A Cat6A cable will need to be run from behind the display to the table box for the touch screen console and content sharing. Should BYOD be required, a second Cat6A should be run from behind the display to the table box."),
                ("Room Options", 
                 "The room has the option of a room booking panel. This is a Teams certified room booking panel, that allows for users to see the status of the room (red for occupied, green for available), book the room from the touch screen, and also book it as a Teams meeting. Room booking panels may have light bars added for easy identification of room availability, and occupancy sensors, which release unoccupied rooms from unused bookings.")
            ]

        # --- 86" ROOM (Up to 6.5m) ---
        elif "86" in r_type:
            return [
                ("Proposed Solution", 
                 "The 10P meeting room represents rooms with a maximum viewing distance of up to 6.5m. Each 10P meeting room shall use an 86” display, with a Maxhub W70 Bar. This Windows 11 based unit has a quad camera, best in class AI based audio and robust cloud based monitoring.\nAt the table, USB-C connectivity is included for content sharing. BYOD is also possible via this cable. It is recommended that MS Teams shall be the primary connectivity method using the inbuilt wireless connectivity."),
                ("Works in Association", 
                 "Behind the LCD, mounted offset to avoid the LCD bracket, there will need to be 1x Double GPO, and two data points (Teams Compute, Display). A Cat6A cable will need to be run from behind the display to the table box for the touch screen console and content sharing. Should BYOD be required, a second Cat6A should be run from behind the display to the table box."),
                ("Room Options", 
                 "The room has the option of a room booking panel. This is a Teams certified room booking panel, that allows for users to see the status of the room (red for occupied, green for available), book the room from the touch screen, and also book it as a Teams meeting. Room booking panels may have light bars added for easy identification of room availability, and occupancy sensors, which release unoccupied rooms from unused bookings.")
            ]

        # --- 98" ROOM (Up to 7.5m) ---
        elif "98" in r_type:
            return [
                ("Proposed Solution", 
                 "The 16P meeting room represents the large meeting room which has a furthest participant of approximately 7.5m. The room shall require a 98” display, with a Maxhub W70 Bar. This Windows 11 based unit has a quad camera, best in class AI based audio and robust cloud based monitoring.\nAt the table, USB-C connectivity is included for content sharing. BYOD is also possible via this cable. It is recommended that MS Teams shall be the primary connectivity method using the inbuilt wireless connectivity.\n\nAt 7.5m, the furthest participant in this room represents the limits of the range of the audio pickup with the bar, even with the AI enhancements. The ambient noise in this room should be minimised, and an RT60 value of less than 0.5 seconds achieved. Glass on either side of the room can cause acoustic challenges, and should the acoustic conditions not be able to be guaranteed, we recommend the expanded audio option"),
                ("Audio Option", 
                 "Should the room have acoustic difficulties, we recommend a dedicated audio system in the room. A central Sennheiser ceiling microphone brings all participants within 3m of a microphone element, and when combined with the audio processor allows us to tune the room to overcome the acoustic challenges. Ceiling speakers are included to cover the room and be tuned in conjunction with the microphone. All audio equipment is cloud monitored and supported, with real time AI based audio tuning ensuring the room sounds as it should even in changing acoustic conditions."),
                ("Further Options", 
                 "The room has the option of a room booking panel. This is a Teams certified room booking panel, that allows for users to see the status of the room (red for occupied, green for available), book the room from the touch screen, and also book it as a Teams meeting. Room booking panels may have light bars added for easy identification of room availability, and occupancy sensors, which release unoccupied rooms from unused bookings."),
                ("Works in Association", 
                 "Behind the LCD, mounted offset to avoid the LCD bracket, there will need to be 1x Double GPO, and two data points (Teams Compute, Display). A Cat6A cable will need to be run from behind the display to the table box for the touch screen console and content sharing. Should BYOD be required, a second Cat6A should be run from behind the display to the table box.\n\nShould the audio option be added, a further double GPO and three data points shall be added behind the display. A data shall be required in the ceiling for the microphone.")
            ]
        
        # --- DEFAULT (For unknown) ---
        else:
             return [
                ("Proposed Solution", "Standard fit-out solution as per Bill of Materials."),
                ("Works in Association", "Standard power and data requirements apply."),
            ]

    def generate_multi_room_proposal(client_name, room_list, project_mode, fitout_pkgs):
        
        # --- 0. SORT ROOM LIST (Smallest to Largest) ---
        room_list.sort(key=lambda x: x['distance'])

        # --- 1. SELECT TEMPLATE ---
        script_dir = os.path.dirname(os.path.abspath(__file__))
        
        if project_mode == "Data#3 (Cisco)":
            template_filename = "Template_Data3.docx"
        else:
            template_filename = "Template_Fitout.docx"

        template_path = os.path.join(script_dir, template_filename)
        
        if os.path.exists(template_path):
            doc = Document(template_path)
            doc.add_page_break() 
        else:
            doc = Document()

        # --- SET FONT TO HELVETICA ---
        try:
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Helvetica'
            font.size = Pt(10)
        except: pass 

        try:
            section = doc.sections[0]
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
        except: pass

        # --- HELPERS ---
        def shade_cell(cell, color_hex):
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
            cell._tc.get_or_add_tcPr().append(shading_elm)

        def add_manual_heading(text, size, color_rgb=None):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Helvetica'
            run.font.size = Pt(size)
            if color_rgb: run.font.color.rgb = color_rgb
            return p
        
        # --- IMPROVED TEXT FORMATTING HELPER (For Body Text) ---
        def add_body_text(text):
            p_body = doc.add_paragraph(text)
            p_body.paragraph_format.space_after = Pt(6) 
            p_body.paragraph_format.line_spacing = 1.15  
            p_body.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p_body.runs:
                run.font.name = 'Helvetica'
                run.font.size = Pt(11) # Size 11
            return p_body

        # --- IMPROVED TEXT FORMATTING HELPER (For Heading + Body Blocks) ---
        def add_bold_heading_text(heading, text):
            # Heading Paragraph
            p_head = doc.add_paragraph()
            p_head.paragraph_format.space_before = Pt(12)
            p_head.paragraph_format.space_after = Pt(2) # Tight to the text below
            run_h = p_head.add_run(heading)
            run_h.bold = True
            run_h.font.name = 'Helvetica'
            run_h.font.size = Pt(11) # Size 11 Heading

            # Body Paragraph using common helper
            add_body_text(text)
            
        def format_row(row, height_cm):
            row.height = Cm(height_cm)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # ---------------------------------------------------------
        # PAGE 1: EXECUTIVE SUMMARY
        # ---------------------------------------------------------
        
        # --- TITLE SECTION REMOVED PER REQUEST ---
        
        add_manual_heading('Partnership Overview', 14)
        
        if project_mode == "Data#3 (Cisco)":
            overview_text = (
                "Alder Technology is pleased to partner with Data#3 to provide this solution. "
                "This document is split into two sections:\n"
                "1. A Master Financial Summary (Hardware + Year 1 Services).\n"
                "2. Detailed Bill of Materials for each specific room.\n\n"
                "Please note: Cisco hardware is listed for engineering reference but is to be supplied and priced by Data#3."
            )
        else:
            overview_text = (
                "Alder Technology is pleased to provide this comprehensive Audio Visual proposal for a complete office fit-out. "
                "This document outlines the Master Financial Summary and the Detailed Bill of Materials for every room, "
                "including all visual displays, conferencing bars, and installation services."
            )
            
        add_body_text(overview_text)

        add_manual_heading('1. Master Room Summary & Pricing', 14)
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        hdr_row = table.rows[0]
        format_row(hdr_row, 1.0)
        hdr = hdr_row.cells
        
        hdr[0].text = "Room Name"
        hdr[1].text = "Classification"
        hdr[2].text = "Supply & Services" 
        # UPDATED HEADER LABEL
        hdr[3].text = "Managed Service P/A\n(5 Years)"
        hdr[4].text = "Total Year 1 (Ex GST)"

        for cell in hdr:
            shade_cell(cell, "D9E2F3")
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        grand_total_project = 0
        
        # COUNTERS FOR CONSOLIDATED OPTIONS
        qty_booking_panels = 0
        qty_audio_upgrades = 0

        # --- CALCULATE TOTALS ---
        for room in room_list:
            name = room['name']
            r_type = room['type'] 
            dist = room['distance'] 

            # --- FIT OUT CALCULATION ---
            if "Fit-Out" in r_type:
                pkg = fitout_pkgs.get(r_type)
                if not pkg and room.get('pkg_key'):
                     pkg = fitout_pkgs.get(room['pkg_key'])

                if not pkg: 
                    continue 

                ms_annual = pkg['ms_price']
                
                # Handle Dual Screen Cost
                qty_display = 2 if "Dual" in r_type else 1
                cost_display = pkg['display'][2] * qty_display
                
                cost_mount = pkg['mount'][2]
                cost_vc = pkg['vc'][2]
                cost_cables = pkg['cables'][2]
                cost_services = pkg['services']
                
                cost_extras = sum([item[2] * item[3] for item in pkg['items']])

                upfront_cost = cost_display + cost_mount + cost_vc + cost_cables + cost_services + cost_extras
                
                display_name = room.get('pkg_key', r_type)
                display_label = f"{display_name} ({dist}m)"
                
                # INCREMENT COUNTERS
                qty_booking_panels += 1 
                if 'audio_upgrade' in pkg:
                    qty_audio_upgrades += 1
                
            # --- DATA#3 CALCULATION ---
            else:
                data = room['config']
                upfront_cost = data['display_price'] + data['mount_price'] + data['cables_price'] + data['service_price']
                ms_annual = data['ms_annual']
                display_label = f"{data['tier_name']} ({dist}m)"

            room_total_y1 = upfront_cost + ms_annual
            grand_total_project += room_total_y1

            row_obj = table.add_row()
            format_row(row_obj, 0.9)
            row = row_obj.cells
            
            row[0].text = name
            row[1].text = display_label
            row[2].text = f"${upfront_cost:,.0f}"
            row[3].text = f"${ms_annual:,.0f}"
            row[4].text = f"${room_total_y1:,.2f}"
            
            for i in range(2, 5):
                row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph("\n")
        p_total = doc.add_paragraph()
        p_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        runner = p_total.add_run(f"TOTAL YEAR 1 PROJECT VALUE (EX GST): ${grand_total_project:,.2f}")
        runner.bold = True
        runner.font.size = Pt(16)
        runner.font.color.rgb = RGBColor(0, 102, 204)

        # --- CONSOLIDATED OPTIONAL UPGRADES TABLE ---
        if qty_booking_panels > 0 or qty_audio_upgrades > 0:
            add_manual_heading('Optional Upgrades (Not included in Total above)', 12, RGBColor(255, 0, 0))
            
            opt_table = doc.add_table(rows=1, cols=5)
            opt_table.style = 'Table Grid'
            # AutoFit enabled (No manual widths)

            h_row = opt_table.rows[0]
            format_row(h_row, 0.9)
            h_cells = h_row.cells
            h_cells[0].text = "Upgrade Item"
            h_cells[1].text = "Qty"
            h_cells[2].text = "Description"
            h_cells[3].text = "Unit Cost"
            h_cells[4].text = "Total Cost"
            for c in h_cells:
                shade_cell(c, "E7E6E6")
                c.paragraphs[0].runs[0].bold = True

            # ROW 1: BOOKING PANELS
            if qty_booking_panels > 0:
                price_unit = 2200.00
                total_panel = price_unit * qty_booking_panels
                
                r_bp = opt_table.add_row()
                format_row(r_bp, 0.8)
                cells_bp = r_bp.cells
                cells_bp[0].text = "Room Booking Panel"
                cells_bp[1].text = str(qty_booking_panels)
                cells_bp[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cells_bp[2].text = "Crestron TS-1070 with Lightbar kit and Multi Surface Mount"
                cells_bp[3].text = f"${price_unit:,.2f}"
                cells_bp[4].text = f"${total_panel:,.2f}"
                cells_bp[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cells_bp[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # ROW 2: AUDIO UPGRADES (IF ANY 98" ROOMS)
            if qty_audio_upgrades > 0:
                pkg_98 = fitout_pkgs.get("Fit-Out 98")
                if pkg_98 and 'audio_upgrade' in pkg_98:
                    upg_data = pkg_98['audio_upgrade']
                    unit_price = upg_data['total_price']
                    total_upg = unit_price * qty_audio_upgrades

                    r_au = opt_table.add_row()
                    format_row(r_au, 0.8)
                    cells_au = r_au.cells
                    cells_au[0].text = "Premium Audio"
                    cells_au[1].text = str(qty_audio_upgrades)
                    cells_au[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cells_au[2].text = upg_data['name']
                    cells_au[3].text = f"${unit_price:,.2f}"
                    cells_au[4].text = f"${total_upg:,.2f}"
                    cells_au[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cells_au[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # ---------------------------------------------------------
        # PAGE 2+: DETAILED ROOM BREAKDOWNS
        # ---------------------------------------------------------
        
        doc.add_page_break()
        add_manual_heading('2. Detailed Room Specifications', 16)

        for room in room_list:
            name = room['name']
            r_type = room['type']
            dist = room['distance']

            # --- HEADER ---
            clean_type_name = room.get('pkg_key', r_type) if "Fit-Out" in r_type else r_type

            # Create a small table for the header (Blue bar)
            table_hdr = doc.add_table(rows=1, cols=1)
            table_hdr.style = 'Table Grid'
            table_hdr.autofit = False
            table_hdr.columns[0].width = Cm(18.0) # Full width
            
            row_hdr = table_hdr.rows[0]
            format_row(row_hdr, 1.0)
            cell_hdr = row_hdr.cells[0]
            cell_hdr.text = f"ROOM: {name} - {clean_type_name} (Furthest Participant: {dist}m)"
            shade_cell(cell_hdr, "1F4E79")
            cell_hdr.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell_hdr.paragraphs[0].runs[0].bold = True

            # --- IMAGE PLACEHOLDER ---
            table_img = doc.add_table(rows=1, cols=1)
            table_img.style = 'Table Grid'
            table_img.autofit = False
            table_img.columns[0].width = Cm(18.0)
            
            row_img = table_img.rows[0]
            format_row(row_img, 4.0)
            cell_img = row_img.cells[0]
            cell_img.text = "[PASTE FLOOR PLAN IMAGE HERE]"
            cell_img.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # --- INSERT TEXT BLOCKS ---
            if "Fit-Out" in r_type:
                text_blocks = get_fitout_text_blocks(clean_type_name)
                for heading, body in text_blocks:
                    add_bold_heading_text(heading, body)

            # --- BOM TABLE (3 Cols, AutoFit to 18cm) ---
            table_room = doc.add_table(rows=0, cols=3)
            table_room.style = 'Table Grid'
            table_room.autofit = False 
            table_room.columns[0].width = Cm(4.0) # Item
            table_room.columns[1].width = Cm(1.5) # Qty
            table_room.columns[2].width = Cm(12.5) # Description

            # --- COL HEADERS ---
            row_cols_obj = table_room.add_row()
            format_row(row_cols_obj, 0.9)
            row_cols = row_cols_obj.cells
            row_cols[0].text = "Item"
            row_cols[1].text = "Qty"
            row_cols[2].text = "Description / Model"
            
            for c in row_cols:
                shade_cell(c, "D9E2F3")
                c.paragraphs[0].runs[0].bold = True

            # =========================================================
            # MODE A: FIT-OUT LOGIC
            # =========================================================
            if "Fit-Out" in r_type:
                pkg = fitout_pkgs.get(r_type)
                if not pkg and room.get('pkg_key'): pkg = fitout_pkgs.get(room['pkg_key'])

                row_sec1_obj = table_room.add_row()
                format_row(row_sec1_obj, 0.8)
                row_sec1 = row_sec1_obj.cells
                row_sec1[0].merge(row_sec1[2])
                row_sec1[0].text = "1. Hardware & Services Scope"
                shade_cell(row_sec1[0], "E7E6E6")
                row_sec1[0].paragraphs[0].runs[0].bold = True

                def add_row(cat, qty, desc):
                    r_obj = table_room.add_row()
                    format_row(r_obj, 0.9)
                    r = r_obj.cells
                    r[0].text = cat
                    r[1].text = str(qty)
                    r[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r[2].text = desc

                # 1. VC BAR
                add_row(pkg['vc'][0], 1, pkg['vc'][1])
                
                # 2. DISPLAY
                disp_qty = 2 if "Dual" in r_type else 1
                add_row(pkg['display'][0], disp_qty, pkg['display'][1])
                
                # 3. EXTRA ITEMS
                for item in pkg['items']:
                    add_row(item[0], item[3], item[1])
                
                # 4. MOUNT & CABLES
                add_row(pkg['mount'][0], 1, pkg['mount'][1])
                add_row(pkg['cables'][0], 1, pkg['cables'][1])

                # 6. SERVICES (FIXED)
                r_svc_obj = table_room.add_row()
                format_row(r_svc_obj, 0.9)
                r_svc = r_svc_obj.cells
                r_svc[0].text = "Services"
                r_svc[1].text = "1"
                r_svc[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER # Center Qty
                r_svc[2].text = "Total Services (Staging, Installation, PM, Engineering)"

                # --- MANAGED SERVICES ---
                row_sec3_obj = table_room.add_row()
                format_row(row_sec3_obj, 0.8)
                row_sec3 = row_sec3_obj.cells
                row_sec3[0].merge(row_sec3[2])
                row_sec3[0].text = "2. Managed Services"
                shade_cell(row_sec3[0], "E7E6E6")
                row_sec3[0].paragraphs[0].runs[0].bold = True

                r_msa_obj = table_room.add_row()
                format_row(r_msa_obj, 0.9)
                r_msa = r_msa_obj.cells
                r_msa[0].text = "Support"
                r_msa[1].text = "1"
                r_msa[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER # Center Qty
                r_msa[2].text = "Managed Service Agreement - Year 1 (Annual Billing)"

                # --- 4. AUDIO UPGRADE TABLE (Attached to Room if 98") ---
                if 'audio_upgrade' in pkg:
                    upg = pkg['audio_upgrade']
                    # Separate header for upgrade
                    r_upg_h = table_room.add_row()
                    format_row(r_upg_h, 0.8)
                    c_upg = r_upg_h.cells
                    c_upg[0].merge(c_upg[2])
                    c_upg[0].text = "3. Optional Upgrade: " + upg['name']
                    shade_cell(c_upg[0], "FCE4D6") # Distinctive orange/peach color
                    c_upg[0].paragraphs[0].runs[0].bold = True
                    
                    # Note: No price shown here per user request
                    for item_code, qty, desc, price in upg['items']:
                        add_row(item_code, qty, desc)

            # =========================================================
            # MODE B: DATA#3 (CISCO) LOGIC
            # =========================================================
            else:
                data = room['config']
                
                final_cisco_list = []
                moved_to_alder_list = []

                for item in data['cisco_items']:
                    if not item: continue
                    clean_item = item.strip()
                    if "Shure" in clean_item:
                        moved_to_alder_list.append(clean_item)
                    else:
                        item_desc = clean_item
                        if "Room Bar Pro" in clean_item: item_desc += " / CS-BARPRO-K9"
                        elif "Room Bar" in clean_item and "Pro" not in clean_item: item_desc += " / CS-BAR-T-C-K9"
                        elif "Ceiling Microphone Pro" in clean_item: item_desc += " / CS-MIC-CLGPRO="
                        elif "Wire Hanging" in clean_item: item_desc += " / CS-MIC-CLGP-WHK="
                        final_cisco_list.append(item_desc)

                if final_cisco_list:
                    row_sec1_obj = table_room.add_row()
                    format_row(row_sec1_obj, 0.8)
                    row_sec1 = row_sec1_obj.cells
                    row_sec1[0].merge(row_sec1[2])
                    row_sec1[0].text = "1. Data#3 Supply Scope (Cisco Hardware)"
                    shade_cell(row_sec1[0], "E7E6E6")
                    row_sec1[0].paragraphs[0].runs[0].bold = True

                    for item in final_cisco_list:
                        r_obj = table_room.add_row()
                        format_row(r_obj, 0.9)
                        r = r_obj.cells
                        for cell in r: shade_cell(cell, "FFF2CC")
                        r[0].text = "Video Conf"
                        r[1].text = "1"
                        r[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r[2].text = item 

                row_sec2_obj = table_room.add_row()
                format_row(row_sec2_obj, 0.8)
                row_sec2 = row_sec2_obj.cells
                row_sec2[0].merge(row_sec2[2])
                row_sec2[0].text = "2. Alder Technology Supply Scope"
                shade_cell(row_sec2[0], "E7E6E6")
                row_sec2[0].paragraphs[0].runs[0].bold = True

                def add_spec_row(item_cat, qty, desc_text):
                    r_obj = table_room.add_row()
                    format_row(r_obj, 0.9)
                    r = r_obj.cells
                    r[0].text = item_cat
                    r[1].text = str(qty)
                    r[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r[2].text = desc_text

                add_spec_row("Visual Display", 1, data['display_model'])
                add_spec_row("Mounting", 1, data['mount_model'])
                for moved_item in moved_to_alder_list:
                    add_spec_row("Conf/Audio", 1, moved_item) 
                add_spec_row("Cabling", 1, data['cables_misc'])
                
                r_svc_obj = table_room.add_row()
                format_row(r_svc_obj, 0.9)
                r_svc = r_svc_obj.cells
                r_svc[0].text = "Services"
                r_svc[1].text = "1"
                r_svc[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER # Center Qty
                r_svc[2].text = "Professional Services: Installation, Staging & PM"

                row_sec3_obj = table_room.add_row()
                format_row(row_sec3_obj, 0.8)
                row_sec3 = row_sec3_obj.cells
                row_sec3[0].merge(row_sec3[2])
                row_sec3[0].text = "3. Managed Services"
                shade_cell(row_sec3[0], "E7E6E6")
                row_sec3[0].paragraphs[0].runs[0].bold = True

                r_msa_obj = table_room.add_row()
                format_row(r_msa_obj, 0.9)
                r_msa = r_msa_obj.cells
                r_msa[0].text = "Support"
                r_msa[1].text = "1"
                r_msa[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER # Center Qty
                r_msa[2].text = "Managed Service Agreement - Year 1 (Annual Billing)"

            doc.add_paragraph("")

        # ---------------------------------------------------------
        # FINAL SECTION
        # ---------------------------------------------------------
        doc.add_page_break()
        add_manual_heading('Managed Service Agreement', 14)
        
        msa_text = (
            "Pricing excludes GST and is charged annually with an increase each year of 4% or CPI whichever is the greater. "
            "Acceptance of a 60 month agreement upfront locks pricing for the five year term with no increase for CPI. "
            "Pricing includes all cloud monitoring hosting charges and any onsite support required."
        )
        add_body_text(msa_text)

        add_manual_heading('Exclusions', 14)
        
        # --- EXCLUSIONS TEXT (Replaces Bullet Points) ---
        exclusions_text = (
            "We exclude all power and data, plus building works. All works required shall be identified and must be completed prior to our installers attending site.\n"
            "Out of hours work is NOT included in this proposal.\n"
            "We exclude all height access equipment, and all furniture protection equipment/coverings.\n"
            "All Teams/Exchange credentials must be provided prior to attending site. Admin credentials for any Teams endpoint must be provided to Alder Technology for the duration of any deployment including temporary Teams administrator cloud tenant access and other software admin access for the duration of deployment.\n"
            "The network must be active and configured for Teams prior to attending site.\n"
            "All external services provided by the client or their nominated systems integrator, including Teams Room accounts, Azure Intune registrations and specific deployment requirements, Exchange, Skype for Business and Microsoft security and compliance requirements and Fast track or Peering ISP plans are the responsibility of the client.\n"
            "Microsoft updates and the impact on the hardware, user experience and the operational state of the system are the sole responsibility of the client. Any such updates that require Alder Technology site attendance incur a Service call out fee and hourly rate at agreed hourly rates, unless a service level agreement covering these works is in place.\n"
            "A site planner shall be provided by Alder Technology and must be completed by the client prior to our installers attending site.\n"
            "Should works not be able to commence due to the above or client led delays, a call out fee may be applied.\n"
            "Should a managed service not be engaged, a DLP period of three (3) months is applicable.\n"
            "Alder Technology works with partners to deliver the best product possible. Alder Technology takes full responsibility for all parties and provides a single point of contact and management.\n"
            "Quotes are valid for 30 days unless otherwise specified. Delays or pauses in works due to client delays or room unavailability may result in additional charges."
        )
        
        add_body_text(exclusions_text)

        doc.add_paragraph("\nThank you for your consideration. Please call me if you have any further queries.")
        doc.add_paragraph("Regards,")
        sig = doc.add_paragraph("George Coles")
        sig.runs[0].bold = True

        desktop = os.path.expanduser("~/Desktop")
        save_folder = os.path.join(desktop, "Alder_Quotes")
        if not os.path.exists(save_folder):
            os.makedirs(save_folder)

        safe_client_name = "".join([c for c in client_name if c.isalpha() or c.isdigit() or c==' ']).rstrip()
        timestamp = datetime.now().strftime("%H-%M-%S")
        filename = f"Alder_Quote_{safe_client_name}_{project_mode[:4]}_{timestamp}.docx"
        full_path = os.path.join(save_folder, filename)
        
        doc.save(full_path)
        return full_path

    # ==========================================
    # PART 2: THE USER INTERFACE (GUI)
    # ==========================================

    ctk.set_appearance_mode("Light") 
    ctk.set_default_color_theme("green") 

    app = ctk.CTk()
    app.geometry("1200x800") 
    app.title("Alder Technology - Quoting Tool")
    
    # Attempt to maximize window (Platform dependent)
    try:
        app.state('zoomed')
    except:
        app.after(0, lambda: app.wm_state('zoomed'))

    # Enable HighDPI scaling
    ctk.set_widget_scaling(1.0) 
    
    PRICELIST_DATA, FITOUT_PACKAGES_DYN, EXCEL_ERROR = load_internal_data()
    ADDED_ROOMS = [] 
    DROPDOWN_MAPPING = {}

    # --- ASSETS ---
    script_dir = os.path.dirname(os.path.abspath(__file__))
    logo_path = os.path.join(script_dir, "alder_logo.png")
    logo_image = None
    if os.path.exists(logo_path):
        try:
            pil_image = Image.open(logo_path)
            logo_image = ctk.CTkImage(light_image=pil_image, dark_image=pil_image, size=(280, 100)) 
        except: pass

    # --- FUNCTIONS ---
    def refresh_room_list():
        for widget in scroll_frame.winfo_children():
            widget.destroy()

        if not ADDED_ROOMS:
            ctk.CTkLabel(scroll_frame, text="No rooms added yet.", text_color="gray", font=("Arial", 14)).pack(pady=40)
            return

        for index, room in enumerate(ADDED_ROOMS):
            # Card Frame
            card = ctk.CTkFrame(scroll_frame, fg_color="white", border_width=1, border_color="#E0E0E0", corner_radius=8)
            card.pack(fill="x", pady=5, padx=5)

            # Icon/Type
            if "Fit-Out" in room['type']:
                icon = "📦"
                type_color = "#1F4E79"
            else:
                icon = "🔌"
                type_color = "#009A44"

            # Content Layout
            # Left: Name & Type
            info_frame = ctk.CTkFrame(card, fg_color="transparent")
            info_frame.pack(side="left", padx=15, pady=10)
            
            ctk.CTkLabel(info_frame, text=room['name'], font=("Arial", 16, "bold"), text_color="#333").pack(anchor="w")
            ctk.CTkLabel(info_frame, text=f"{icon} {room['type']}", font=("Arial", 12), text_color=type_color).pack(anchor="w")

            # Right: Delete Button
            btn_del = ctk.CTkButton(
                card, text="Remove", width=80, height=30,
                fg_color="#FFEEEE", text_color="#FF5555", hover_color="#FFDDDD",
                font=("Arial", 12, "bold"),
                command=lambda i=index: delete_room(i)
            )
            btn_del.pack(side="right", padx=15)

        # Update Count
        lbl_count.configure(text=f"Total Rooms: {len(ADDED_ROOMS)}")

    def delete_room(index):
        del ADDED_ROOMS[index]
        refresh_room_list()
    
    def on_dropdown_change(selected_val):
        if "(" in selected_val:
            suggested = selected_val.split("(")[0].strip()
        else:
            suggested = selected_val
        current = entry_room_name.get()
        if not current:
            entry_room_name.insert(0, suggested)

    def update_dropdown_options(choice):
        global DROPDOWN_MAPPING
        DROPDOWN_MAPPING = {}
        display_options = []

        if choice == "Data#3 (Cisco)":
            sorted_tiers = sorted([t for t in PRICELIST_DATA], key=lambda x: x['max_distance'])
            prev_dist = 0
            for tier in sorted_tiers:
                curr = tier['max_distance']
                label = f"{tier['tier_name']} ({prev_dist}m - {curr}m)"
                display_options.append(label)
                DROPDOWN_MAPPING[label] = tier
                prev_dist = curr
        else:
            if FITOUT_PACKAGES_DYN:
                pkg_list = []
                for k, v in FITOUT_PACKAGES_DYN.items():
                    pkg_list.append((k, v))
                pkg_list.sort(key=lambda x: x[1]['max_distance'])
                prev_dist = 0
                for name, data in pkg_list:
                    curr = data['max_distance']
                    label = f"{name} ({prev_dist}m - {curr}m)"
                    display_options.append(label)
                    DROPDOWN_MAPPING[label] = name
                    prev_dist = curr
            else:
                display_options = ["No Packages"]
            
        dropdown_type.configure(values=display_options)
        dropdown_type.set(display_options[0] if display_options else "")
        if display_options: on_dropdown_change(display_options[0])

    def on_add_room():
        r_name = entry_room_name.get().strip()
        label = dropdown_type.get()
        mode = dropdown_project_mode.get()

        if not r_name:
            status_bar.configure(text="Error: Missing Room Name", text_color="#FF5555")
            return

        mapped = DROPDOWN_MAPPING.get(label)
        if not mapped: return
        
        if mode == "Data#3 (Cisco)":
             dist = mapped['max_distance']
        else:
             pkg = FITOUT_PACKAGES_DYN.get(mapped)
             dist = pkg['max_distance'] if pkg else 0.0

        room_entry = {'name': r_name, 'distance': dist}

        if mode == "Data#3 (Cisco)":
            room_entry['type'] = mapped['tier_name']
            room_entry['config'] = mapped
        else:
            room_entry['type'] = label
            room_entry['pkg_key'] = mapped
            room_entry['config'] = None
        
        ADDED_ROOMS.append(room_entry)
        entry_room_name.delete(0, "end")
        on_dropdown_change(dropdown_type.get())
        refresh_room_list()
        status_bar.configure(text="Room Added Successfully", text_color="green")

    def on_generate_click():
        status_bar.configure(text="Generating Proposal...", text_color="blue")
        app.update()
        client = entry_client.get().strip()
        if not client:
            status_bar.configure(text="Error: Enter Client Name", text_color="red")
            return
        if not ADDED_ROOMS:
            status_bar.configure(text="Error: No rooms added", text_color="red")
            return

        try:
            fp = generate_multi_room_proposal(client, ADDED_ROOMS, dropdown_project_mode.get(), FITOUT_PACKAGES_DYN)
            status_bar.configure(text=f"Success! Saved to Desktop/Alder_Quotes", text_color="#009A44")
            try: os.startfile(os.path.dirname(fp))
            except: pass
        except Exception as e:
            status_bar.configure(text=f"Error: {str(e)}", text_color="red")

    # --- LAYOUT CONSTRUCTION (GRID) ---
    app.grid_columnconfigure(0, weight=0, minsize=350) # Sidebar
    app.grid_columnconfigure(1, weight=1) # Main Content
    app.grid_rowconfigure(0, weight=1) # Full Height

    # 1. SIDEBAR (Left)
    sidebar = ctk.CTkFrame(app, fg_color="#F9F9F9", corner_radius=0, width=350)
    sidebar.grid(row=0, column=0, sticky="nsew")
    
    # Logo Area
    if logo_image:
        ctk.CTkLabel(sidebar, image=logo_image, text="").pack(pady=(30, 20))
    else:
        ctk.CTkLabel(sidebar, text="ALDER TECH", font=("Arial Black", 24), text_color="#009A44").pack(pady=(40, 20))

    # Controls Container
    ctrl_frame = ctk.CTkFrame(sidebar, fg_color="transparent")
    ctrl_frame.pack(fill="x", padx=20)

    ctk.CTkLabel(ctrl_frame, text="Client Details", font=("Arial", 14, "bold"), text_color="#555").pack(anchor="w", pady=(10, 5))
    entry_client = ctk.CTkEntry(ctrl_frame, placeholder_text="Client Name", height=40, font=("Arial", 14))
    entry_client.pack(fill="x", pady=5)

    ctk.CTkLabel(ctrl_frame, text="Project Scope", font=("Arial", 14, "bold"), text_color="#555").pack(anchor="w", pady=(20, 5))
    dropdown_project_mode = ctk.CTkOptionMenu(ctrl_frame, values=["Data#3 (Cisco)", "Fit-Out (Full Scope)"], height=40, font=("Arial", 13), fg_color="#333", command=update_dropdown_options)
    dropdown_project_mode.pack(fill="x", pady=5)

    ctk.CTkFrame(ctrl_frame, height=2, fg_color="#E0E0E0").pack(fill="x", pady=30) # Divider

    ctk.CTkLabel(ctrl_frame, text="Room Builder", font=("Arial", 16, "bold"), text_color="#009A44").pack(anchor="w", pady=(0, 10))
    
    dropdown_type = ctk.CTkOptionMenu(ctrl_frame, values=[], height=40, fg_color="#1F4E79", font=("Arial", 13), command=on_dropdown_change)
    dropdown_type.pack(fill="x", pady=5)
    
    entry_room_name = ctk.CTkEntry(ctrl_frame, placeholder_text="Room Name", height=40, font=("Arial", 14))
    entry_room_name.pack(fill="x", pady=5)

    btn_add = ctk.CTkButton(ctrl_frame, text="+ ADD ROOM", height=50, fg_color="#009A44", hover_color="#007a36", font=("Arial", 14, "bold"), command=on_add_room)
    btn_add.pack(fill="x", pady=20)

    # 2. MAIN DASHBOARD (Right)
    main_area = ctk.CTkFrame(app, fg_color="white", corner_radius=0)
    main_area.grid(row=0, column=1, sticky="nsew", padx=0, pady=0)
    
    # Header
    top_bar = ctk.CTkFrame(main_area, height=80, fg_color="white")
    top_bar.pack(fill="x", side="top")
    ctk.CTkLabel(top_bar, text="Project Overview", font=("Arial", 28, "bold"), text_color="#333").pack(side="left", padx=40, pady=30)
    lbl_count = ctk.CTkLabel(top_bar, text="Total Rooms: 0", font=("Arial", 14), text_color="gray")
    lbl_count.pack(side="right", padx=40, pady=30)

    # List Area
    scroll_frame = ctk.CTkScrollableFrame(main_area, fg_color="#F4F4F4", corner_radius=0)
    scroll_frame.pack(fill="both", expand=True, padx=40, pady=(0, 20))

    # Bottom Action Area
    action_bar = ctk.CTkFrame(main_area, height=100, fg_color="white")
    action_bar.pack(fill="x", side="bottom", pady=0)
    
    status_bar = ctk.CTkLabel(action_bar, text="Ready", font=("Arial", 12), text_color="gray", anchor="w")
    status_bar.pack(side="left", padx=40)

    btn_gen = ctk.CTkButton(action_bar, text="GENERATE PROPOSAL", width=250, height=55, fg_color="#1F4E79", font=("Arial", 15, "bold"), command=on_generate_click)
    btn_gen.pack(side="right", padx=40, pady=20)

    # Init
    update_dropdown_options("Data#3 (Cisco)")
    refresh_room_list()

    app.mainloop()

except Exception as e:
    root = tkinter.Tk()
    root.withdraw()
    msg = str(traceback.format_exc())
    tkinter.messagebox.showerror("Fatal Error", msg)
