import traceback
import tkinter
import tkinter.messagebox

# --- CRASH REPORTER WRAPPER ---
try:
    import customtkinter as ctk
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from datetime import datetime
    import os
    import sys

    # ==========================================
    # PART 1: THE LOGIC & PRICING DATABASE
    # ==========================================

    def get_room_configuration(distance):
        # --- TIER 1: UP TO 3.0M ---
        if distance <= 3.0:
            return {
                "tier_name": "Small Meeting Space",
                "cisco_items": ["Cisco Room Bar", "Integrated Camera", "Integrated Audio"],
                "display_model": "LG 55UL3J-B (55\" 4K UHD, 400nits, WebOS)",
                "display_price": 1100.00,
                "mount_model": "Venturi VP-F80 (Suit 50\"-75\")",
                "mount_price": 65.00,
                "cables_misc": "HDMI & Patch Leads", 
                "cables_price": 50.00,
                "service_price": 3000.00,
                "ms_annual": 1200.00
            }
        # --- TIER 2: UP TO 4.5M ---
        elif distance <= 4.5:
            return {
                "tier_name": "Medium Meeting Space",
                "cisco_items": ["Cisco Room Bar", "Integrated Camera", "1x Cisco Table Microphone Pro"],
                "display_model": "LG 65UL3J-B (65\" 4K UHD, 400nits, WebOS)",
                "display_price": 1500.00,
                "mount_model": "Venturi VP-F80 (Suit 50\"-75\")",
                "mount_price": 65.00,
                "cables_misc": "HDMI & Patch Leads",
                "cables_price": 50.00,
                "service_price": 3000.00,
                "ms_annual": 1200.00
            }
        # --- TIER 3: UP TO 5.5M ---
        elif distance <= 5.5:
            return {
                "tier_name": "Large Meeting Space",
                "cisco_items": ["Cisco Room Bar Pro", "Integrated Dual Camera", "1x Cisco Ceiling Microphone Pro"],
                "display_model": "LG 75UL3J-B (75\" 4K UHD, 400nits, WebOS)",
                "display_price": 2200.00,
                "mount_model": "Venturi VP-F80 (Suit 50\"-75\")",
                "mount_price": 65.00,
                "cables_misc": "HDMI, Patch Leads & Mount Fixings",
                "cables_price": 80.00,
                "service_price": 3000.00,
                "ms_annual": 1500.00
            }
        # --- TIER 4: UP TO 6.5M ---
        elif distance <= 6.5:
            return {
                "tier_name": "Extra Large Space",
                "cisco_items": ["Cisco Room Bar Pro", "Integrated Dual Camera", "1x Cisco Ceiling Microphone Pro"],
                "display_model": "LG 86UL3J-B (86\" 4K UHD, 330nits, WebOS)",
                "display_price": 3300.00,
                "mount_model": "VP-F100 (Suit 82\"-98\")",
                "mount_price": 100.00,
                "cables_misc": "HDMI, Patch Leads & Heavy Duty Fixings",
                "cables_price": 100.00,
                "service_price": 3500.00,
                "ms_annual": 1500.00
            }
        # --- TIER 5: UP TO 7.5M ---
        elif distance <= 7.5:
            return {
                "tier_name": "Boardroom",
                "cisco_items": ["Cisco Kit EQ + AV Integrator License", "Cisco Quad Cam", "2x Cisco Ceiling Mic Pro", "6-8x Shure Ceiling Speakers"],
                "display_model": "LG 98UM5K (98\" 4K UHD, 500nits)",
                "display_price": 7500.00,
                "mount_model": "VP-F100 (Suit 82\"-98\")",
                "mount_price": 100.00,
                "cables_misc": "HDMI, Patch Leads, Speaker Cable, Fixings",
                "cables_price": 200.00,
                "service_price": 4000.00,
                "ms_annual": 3000.00
            }
        else:
            return None

    def generate_multi_room_proposal(client_name, room_list):
        
        # --- TEMPLATE LOADING ---
        script_dir = os.path.dirname(os.path.abspath(__file__))
        template_path = os.path.join(script_dir, "Alder_Template.docx")
        
        if os.path.exists(template_path):
            doc = Document(template_path)
            doc.add_page_break() 
        else:
            print("Template not found. Creating blank.")
            doc = Document()

        # --- SET FONT TO HELVETICA ---
        try:
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Helvetica'
            font.size = Pt(10)
        except:
            pass 

        # Style margins
        try:
            section = doc.sections[0]
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
        except: pass

        # --- HELPER: SHADE TABLE CELLS ---
        def shade_cell(cell, color_hex):
            """Apply background color to a table cell."""
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
            cell._tc.get_or_add_tcPr().append(shading_elm)

        # --- HELPER: BOLD HEADING ---
        def add_bold_heading(text, level):
            h = doc.add_heading(text, level)
            for run in h.runs:
                run.bold = True
            return h

        # ---------------------------------------------------------
        # PAGE 1: EXECUTIVE SUMMARY & MASTER LIST
        # ---------------------------------------------------------
        
        header = add_bold_heading(f'AV Proposal: {client_name}', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f'Prepared by: Alder Technology')
        doc.add_paragraph(f'Date: {datetime.now().strftime("%d/%m/%Y")}')
        doc.add_paragraph(f'Total Rooms Scoped: {len(room_list)}')
        doc.add_paragraph('------------------------------------------------------')

        add_bold_heading('Partnership Overview', level=2)
        doc.add_paragraph(
            "Alder Technology is pleased to partner with Data#3 to provide this solution. "
            "This document is split into two sections:\n"
            "1. A Master Financial Summary (Hardware + Year 1 Services).\n"
            "2. Detailed Bill of Materials for each specific room.\n\n"
            "Please note: Cisco hardware is listed for engineering reference but is to be supplied and priced by Data#3."
        )

        add_bold_heading('1. Master Room Summary & Pricing', level=2)
        
        # Master Table Columns: Room | Class | Supply & Services | MSA | Total
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        hdr = table.rows[0].cells
        hdr[0].text = "Room Name"
        hdr[1].text = "Classification"
        hdr[2].text = "Supply & Services"  # Removed UPFRONT
        hdr[3].text = "Managed Service (Year 1)"
        hdr[4].text = "Total Year 1 (Ex GST)"

        # Style Header
        for cell in hdr:
            shade_cell(cell, "D9E2F3") # Light Blue
            cell.paragraphs[0].runs[0].bold = True

        grand_total_project = 0

        for room in room_list:
            name = room['name']
            dist = room['distance']
            data = room['config']
            
            upfront_cost = data['display_price'] + data['mount_price'] + data['cables_price'] + data['service_price']
            ms_annual = data['ms_annual'] 
            
            room_total_y1 = upfront_cost + ms_annual
            grand_total_project += room_total_y1

            row = table.add_row().cells
            row[0].text = name
            row[1].text = f"{data['tier_name']} ({dist}m)"
            row[2].text = f"${upfront_cost:,.0f}"
            row[3].text = f"${ms_annual:,.0f}"
            row[4].text = f"${room_total_y1:,.2f}"
            
            # Right Align Numbers
            for i in range(2, 5):
                row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph("\n")
        p_total = doc.add_paragraph()
        p_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        runner = p_total.add_run(f"TOTAL YEAR 1 PROJECT VALUE (EX GST): ${grand_total_project:,.2f}")
        runner.bold = True
        runner.font.size = Pt(16)
        runner.font.color.rgb = RGBColor(0, 102, 204)
        
        doc.add_paragraph("(Includes Alder Hardware, Installation Services, and 1st Year Managed Service)")
        
        # ---------------------------------------------------------
        # PAGE 2+: DETAILED ROOM BREAKDOWNS (6 COLUMNS)
        # ---------------------------------------------------------
        
        doc.add_page_break()
        add_bold_heading('2. Detailed Room Specifications', level=1)
        doc.add_paragraph("The following section details the specific technology and pricing for each room.")

        for room in room_list:
            name = room['name']
            data = room['config']
            
            # --- CREATE ONE UNIFIED TABLE PER ROOM ---
            # Columns: Item | Qty | Model | Description | Unit Cost | Total
            table_room = doc.add_table(rows=0, cols=6)
            table_room.style = 'Table Grid'
            table_room.autofit = True
            
            # Set Column Widths (Approximate)
            # Item(2cm), Qty(1cm), Model(4cm), Desc(5cm), Cost(2.5cm), Total(2.5cm)
            
            # 1. ROOM HEADER ROW (Merges all 6 cols)
            row_hdr = table_room.add_row().cells
            row_hdr[0].merge(row_hdr[1]).merge(row_hdr[2]).merge(row_hdr[3]).merge(row_hdr[4]).merge(row_hdr[5])
            row_hdr[0].text = f"ROOM: {name} ({room['distance']}m) - {data['tier_name']}"
            shade_cell(row_hdr[0], "1F4E79") # Dark Blue
            run = row_hdr[0].paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255) # White text
            run.font.size = Pt(11)

            # 2. COLUMNS HEADER
            row_cols = table_room.add_row().cells
            row_cols[0].text = "Item"
            row_cols[1].text = "Qty"
            row_cols[2].text = "Model"
            row_cols[3].text = "Description"
            row_cols[4].text = "Unit Cost"
            row_cols[5].text = "Total"
            
            for c in row_cols:
                shade_cell(c, "D9E2F3") # Light Blue
                c.paragraphs[0].runs[0].bold = True
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 3. SECTION: CISCO (DATA#3)
            row_sec1 = table_room.add_row().cells
            row_sec1[0].merge(row_sec1[5]) # Merge across
            row_sec1[0].text = "1. Data#3 Supply Scope (Cisco Hardware)"
            shade_cell(row_sec1[0], "E7E6E6") # Light Grey
            row_sec1[0].paragraphs[0].runs[0].bold = True

            for item in data['cisco_items']:
                r = table_room.add_row().cells
                r[0].text = "Video Conf"
                r[1].text = "1"
                r[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                r[2].text = item
                r[3].text = "Supplied by Data#3"
                r[4].text = "Excl."
                r[5].text = "Excl."
                r[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # 4. SECTION: ALDER SUPPLY
            row_sec2 = table_room.add_row().cells
            row_sec2[0].merge(row_sec2[5])
            row_sec2[0].text = "2. Alder Technology Supply Scope" # Removed Upfront
            shade_cell(row_sec2[0], "E7E6E6") # Light Grey
            row_sec2[0].paragraphs[0].runs[0].bold = True

            # Helper to add 6-col rows
            def add_spec_row(item_cat, qty, model_txt, price_val):
                r = table_room.add_row().cells
                r[0].text = item_cat
                r[1].text = str(qty)
                r[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Intelligent Split: "LG 55 (4K UHD)" -> Model: LG 55, Desc: (4K UHD)
                if "(" in model_txt:
                    parts = model_txt.split("(", 1)
                    r[2].text = parts[0].strip()
                    r[3].text = "(" + parts[1].strip()
                else:
                    r[2].text = model_txt
                    r[3].text = ""

                r[4].text = f"${price_val:,.2f}"
                r[5].text = f"${(price_val * qty):,.2f}"
                r[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            add_spec_row("Visual Display", 1, data['display_model'], data['display_price'])
            add_spec_row("Mounting", 1, data['mount_model'], data['mount_price'])
            add_spec_row("Cabling", 1, data['cables_misc'], data['cables_price'])
            
            # Services usually don't have a model, so handle manually
            r_svc = table_room.add_row().cells
            r_svc[0].text = "Services"
            r_svc[1].text = "1"
            r_svc[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_svc[2].text = "Professional Services"
            r_svc[3].text = "Installation, Staging & PM"
            r_svc[4].text = f"${data['service_price']:,.2f}"
            r_svc[5].text = f"${data['service_price']:,.2f}"
            r_svc[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_svc[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Subtotal Row
            upfront_subtotal = data['display_price'] + data['mount_price'] + data['cables_price'] + data['service_price']
            row_sub = table_room.add_row().cells
            # Merge 0-4 for label, 5 for price
            row_sub[0].merge(row_sub[1]).merge(row_sub[2]).merge(row_sub[3]).merge(row_sub[4])
            row_sub[0].text = "SUB-TOTAL (EX GST):" # Removed Upfront
            row_sub[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_sub[0].paragraphs[0].runs[0].bold = True
            
            row_sub[5].text = f"${upfront_subtotal:,.2f}"
            row_sub[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_sub[5].paragraphs[0].runs[0].bold = True

            # 5. SECTION: MANAGED SERVICES
            row_sec3 = table_room.add_row().cells
            row_sec3[0].merge(row_sec3[5])
            row_sec3[0].text = "3. Managed Services"
            shade_cell(row_sec3[0], "E7E6E6") # Light Grey
            row_sec3[0].paragraphs[0].runs[0].bold = True

            # MSA Row
            r_msa = table_room.add_row().cells
            r_msa[0].text = "Support"
            r_msa[1].text = "1"
            r_msa[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_msa[2].text = "Managed Service Agreement"
            r_msa[3].text = "Year 1 (Annual Billing)"
            r_msa[4].text = f"${data['ms_annual']:,.2f}"
            r_msa[5].text = f"${data['ms_annual']:,.2f}"
            r_msa[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_msa[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            doc.add_paragraph("") # Tiny gap between tables

        # ---------------------------------------------------------
        # FINAL SECTION: FURTHER OPTIONS & MSA & EXCLUSIONS
        # ---------------------------------------------------------
        
        doc.add_page_break()
        
        # --- FURTHER OPTIONS ---
        add_bold_heading('Further Options', level=2)
        
        room_count = len(room_list)
        panel_cost = 2200.00
        total_panel_cost = room_count * panel_cost
        
        table_opt = doc.add_table(rows=2, cols=4)
        table_opt.style = 'Table Grid'
        
        headers = table_opt.rows[0].cells
        headers[0].text = "Item"
        headers[1].text = "Quantity"
        headers[2].text = "Cost per Room"
        headers[3].text = "Total"
        
        for c in headers: 
            shade_cell(c, "D9E2F3")
            c.paragraphs[0].runs[0].bold = True
        
        row = table_opt.rows[1].cells
        row[0].text = "Room Booking Panel"
        row[1].text = str(room_count)
        row[2].text = f"${panel_cost:,.2f}"
        row[3].text = f"${total_panel_cost:,.2f}"
        
        doc.add_paragraph("\n") 

        # --- MSA ---
        add_bold_heading('Managed Service Agreement', level=2)
        doc.add_paragraph(
            "An example costing is provided below. Please note the below does not constitute a quote until the scope of the Managed Service Agreement is finalised. "
            "Pricing excludes GST and is charged annually with an increase each year of 4% or CPI whichever is the greater. "
            "Acceptance of a 60 month agreement upfront locks pricing for the five year term with no increase for CPI. "
            "Pricing includes all cloud monitoring hosting charges and any onsite support required."
        )

        # --- EXCLUSIONS ---
        add_bold_heading('Exclusions', level=2)
        
        exclusions = [
            "We exclude all power and data, plus building works. All works required shall be identified and must be completed prior to our installers attending site.",
            "We exclude all height access equipment, and all furniture protection equipment/coverings.",
            "All Teams/Exchange credentials must be provided prior to attending site. Admin credentials for any Teams endpoint must be provided to Alder Technology for the duration of any deployment including temporary Teams administrator cloud tenant access and other software admin access for the duration of deployment.",
            "The network must be active and configured for Teams prior to attending site.",
            "All external services provided by the client or their nominated systems integrator, including Teams Room accounts, Azure Intune registrations and specific deployment requirements, Exchange, Skype for Business and Microsoft security and compliance requirements and Fast track or Peering ISP plans are the responsibility of the client.",
            "Microsoft updates and the impact on the hardware, user experience and the operational state of the system are the sole responsibility of the client. Any such updates that require Alder Technology site attendance incur a Service call out fee and hourly rate at agreed hourly rates, unless a service level agreement covering these works is in place.",
            "A site planner shall be provided by Alder Technology and must be completed by the client prior to our installers attending site.",
            "Should works not be able to commence due to the above or client led delays, a call out fee may be applied.",
            "Should a managed service not be engaged, a DLP period of three (3) months is applicable.",
            "Alder Technology works with partners to deliver the best product possible. Alder Technology takes full responsibility for all parties and provides a single point of contact and management."
        ]
        
        for item in exclusions:
            p = doc.add_paragraph(item)
            p.paragraph_format.space_after = Pt(6)

        doc.add_paragraph("\nThank you for your consideration. Please call me if you have any further queries.")
        
        doc.add_paragraph("Regards,")
        sig = doc.add_paragraph("George Coles")
        sig.runs[0].bold = True

        # --- SAVE ---
        desktop = os.path.expanduser("~/Desktop")
        save_folder = os.path.join(desktop, "Alder_Quotes")
        if not os.path.exists(save_folder):
            os.makedirs(save_folder)

        safe_client_name = "".join([c for c in client_name if c.isalpha() or c.isdigit() or c==' ']).rstrip()
        timestamp = datetime.now().strftime("%H-%M-%S")
        filename = f"Alder_Quote_{safe_client_name}_{timestamp}.docx"
        full_path = os.path.join(save_folder, filename)
        
        doc.save(full_path)
        return full_path

    # ==========================================
    # PART 2: THE USER INTERFACE (GUI)
    # ==========================================

    ctk.set_appearance_mode("Dark")
    ctk.set_default_color_theme("blue")

    app = ctk.CTk()
    app.geometry("600x700")
    app.title("Alder Technology - Multi-Room Quoter")

    def on_generate_click():
        lbl_status.configure(text="Processing...", text_color="white")
        
        client = entry_client.get()
        if not client:
            lbl_status.configure(text="Error: Enter Client Name", text_color="#FF5555")
            return

        raw_text = txt_rooms.get("0.0", "end")
        lines = raw_text.split('\n')
        valid_rooms = []
        
        try:
            for line in lines:
                line = line.strip()
                if not line: continue 
                if "\t" in line: line = line.replace("\t", ",") 
                if "," not in line: continue 
                    
                parts = line.split(',')
                r_name = parts[0].strip()
                r_dist_str = parts[-1].lower().replace('m', '').strip()
                
                try:
                    r_dist = float(r_dist_str)
                except ValueError:
                    continue 
                
                config = get_room_configuration(r_dist)
                if config:
                    valid_rooms.append({'name': r_name, 'distance': r_dist, 'config': config})
            
            if len(valid_rooms) == 0:
                lbl_status.configure(text="Error: No valid rooms found.\nFormat: 'Name, Distance'", text_color="#FF5555")
                return

            filepath = generate_multi_room_proposal(client, valid_rooms)
            lbl_status.configure(text=f"SUCCESS!\nSaved to Desktop/Alder_Quotes:\n{os.path.basename(filepath)}", text_color="#00FF00")
            
            try: os.startfile(os.path.dirname(filepath))
            except: pass
            
        except PermissionError:
            lbl_status.configure(text="ERROR: PERMISSION DENIED.\nClose Word and try again.", text_color="#FF5555")
        except Exception as e:
            lbl_status.configure(text=f"Error: {str(e)}", text_color="#FF5555")

    # --- UI ELEMENTS ---
    lbl_title = ctk.CTkLabel(app, text="Alder Multi-Room Quoter", font=("Roboto Medium", 24))
    lbl_title.pack(pady=15)

    entry_client = ctk.CTkEntry(app, placeholder_text="Client Name (e.g. Acme Corp)", width=400)
    entry_client.pack(pady=5)

    lbl_instr = ctk.CTkLabel(app, text="Paste Excel Data below (Name | Depth)", text_color="gray")
    lbl_instr.pack(pady=(15, 5))

    txt_rooms = ctk.CTkTextbox(app, width=400, height=300)
    txt_rooms.pack(pady=5)

    btn_generate = ctk.CTkButton(app, text="GENERATE MASTER PROPOSAL", command=on_generate_click, width=400, height=50, font=("Roboto Medium", 14))
    btn_generate.pack(pady=20)

    # --- TEMPLATE CHECKER ---
    script_dir = os.path.dirname(os.path.abspath(__file__))
    tpl_path = os.path.join(script_dir, "Alder_Template.docx")
    
    if os.path.exists(tpl_path):
        status_text = "Ready (Template Found)"
        status_color = "#00FF00"
    else:
        status_text = f"Ready (Template NOT Found)\nLooking in: {script_dir}"
        status_color = "yellow"

    lbl_status = ctk.CTkLabel(app, text=status_text, text_color=status_color, font=("Roboto", 12))
    lbl_status.pack(pady=10)

    app.mainloop()

# --- CRASH CATCHER ---
except Exception as e:
    root = tkinter.Tk()
    root.withdraw()
    error_msg = str(traceback.format_exc())
    tkinter.messagebox.showerror("Program Error", f"The program crashed:\n\n{error_msg}")
