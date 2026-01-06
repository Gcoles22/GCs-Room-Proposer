import traceback
import tkinter
import tkinter.messagebox
import sys
import os
from datetime import datetime

# --- CRASH REPORTER WRAPPER ---
try:
    import customtkinter as ctk
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from PIL import Image
    import openpyxl  # REQUIRED for Excel reading

    # ==========================================
    # PART 1: LOGIC & EXCEL DATABASE
    # ==========================================

    def load_pricelist_from_excel():
        """
        Reads master_pricelist.xlsx and returns a list of dictionaries.
        """
        script_dir = os.path.dirname(os.path.abspath(__file__))
        file_path = os.path.join(script_dir, "master_pricelist.xlsx")
        
        if not os.path.exists(file_path):
            return None, "File 'master_pricelist.xlsx' not found."

        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            sheet = wb.active
            
            pricelist = []

            # Skip header (row 1), start from row 2
            for row in sheet.iter_rows(min_row=2, values_only=True):
                # Ensure row isn't empty (check Max Distance column)
                if row[0] is None:
                    continue

                # Parse Cisco Items (Split string by comma)
                cisco_raw = row[2] if row[2] else ""
                cisco_list = [x.strip() for x in str(cisco_raw).split(',')]

                entry = {
                    "max_distance": float(row[0]),
                    "tier_name": str(row[1]),
                    "cisco_items": cisco_list,
                    "display_model": str(row[3]),
                    "display_price": float(row[4] or 0),
                    "mount_model": str(row[5]),
                    "mount_price": float(row[6] or 0),
                    "cables_misc": str(row[7]),
                    "cables_price": float(row[8] or 0),
                    "service_price": float(row[9] or 0),
                    "ms_annual": float(row[10] or 0)
                }
                pricelist.append(entry)

            # Sort by distance (ascending) to ensure logic works correctly
            pricelist.sort(key=lambda x: x['max_distance'])
            
            return pricelist, None

        except Exception as e:
            return None, str(e)

    def get_room_configuration(distance, pricelist):
        """
        Compares input distance against the loaded Excel pricelist.
        """
        if not pricelist:
            return None

        # Iterate through sorted tiers. Returns the first tier where 
        # the room distance fits within the tier's max distance.
        for tier in pricelist:
            if distance <= tier['max_distance']:
                return tier
        
        # If distance is larger than the largest tier in Excel
        return None

    def generate_multi_room_proposal(client_name, room_list):
        
        # --- TEMPLATE LOADING ---
        script_dir = os.path.dirname(os.path.abspath(__file__))
        template_path = os.path.join(script_dir, "Alder_Template.docx")
        
        if os.path.exists(template_path):
            doc = Document(template_path)
            doc.add_page_break() 
        else:
            doc = Document()

        # --- SET FONT TO HELVETICA ---
        try:
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Helvetica'
            font.size = Pt(10)
        except:
            pass 

        # Style margins
        try:
            section = doc.sections[0]
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
        except: pass

        # --- HELPERS FOR STYLING ---
        def shade_cell(cell, color_hex):
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
            cell._tc.get_or_add_tcPr().append(shading_elm)

        def add_manual_heading(text, size, color_rgb=None):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Helvetica'
            run.font.size = Pt(size)
            if color_rgb:
                run.font.color.rgb = color_rgb
            return p
            
        def format_row(row, height_cm):
            row.height = Cm(height_cm)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # ---------------------------------------------------------
        # PAGE 1: EXECUTIVE SUMMARY & MASTER LIST
        # ---------------------------------------------------------
        
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_t = p_title.add_run(f'AV Proposal: {client_name}')
        run_t.bold = True
        run_t.font.size = Pt(26)
        
        doc.add_paragraph(f'Prepared by: Alder Technology')
        doc.add_paragraph(f'Date: {datetime.now().strftime("%d/%m/%Y")}')
        doc.add_paragraph(f'Total Rooms Scoped: {len(room_list)}')
        doc.add_paragraph('------------------------------------------------------')

        add_manual_heading('Partnership Overview', 14)
        doc.add_paragraph(
            "Alder Technology is pleased to partner with Data#3 to provide this solution. "
            "This document is split into two sections:\n"
            "1. A Master Financial Summary (Hardware + Year 1 Services).\n"
            "2. Detailed Bill of Materials for each specific room.\n\n"
            "Please note: Cisco hardware is listed for engineering reference but is to be supplied and priced by Data#3."
        )

        add_manual_heading('1. Master Room Summary & Pricing', 14)
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # --- HEADER ROW ---
        hdr_row = table.rows[0]
        format_row(hdr_row, 1.0)
        hdr = hdr_row.cells
        
        hdr[0].text = "Room Name"
        hdr[1].text = "Classification"
        hdr[2].text = "Supply & Services" 
        hdr[3].text = "Managed Service (Year 1)"
        hdr[4].text = "Total Year 1 (Ex GST)"

        for cell in hdr:
            shade_cell(cell, "D9E2F3")
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        grand_total_project = 0

        for room in room_list:
            name = room['name']
            dist = room['distance']
            data = room['config']
            
            upfront_cost = data['display_price'] + data['mount_price'] + data['cables_price'] + data['service_price']
            ms_annual = data['ms_annual'] 
            
            room_total_y1 = upfront_cost + ms_annual
            grand_total_project += room_total_y1

            row_obj = table.add_row()
            format_row(row_obj, 0.9)
            row = row_obj.cells
            
            row[0].text = name
            row[1].text = f"{data['tier_name']} ({dist}m)"
            row[2].text = f"${upfront_cost:,.0f}"
            row[3].text = f"${ms_annual:,.0f}"
            row[4].text = f"${room_total_y1:,.2f}"
            
            for i in range(2, 5):
                row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph("\n")
        p_total = doc.add_paragraph()
        p_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        runner = p_total.add_run(f"TOTAL YEAR 1 PROJECT VALUE (EX GST): ${grand_total_project:,.2f}")
        runner.bold = True
        runner.font.size = Pt(16)
        runner.font.color.rgb = RGBColor(0, 102, 204)
        
        doc.add_paragraph("(Includes Alder Hardware, Installation Services, and 1st Year Managed Service)")
        
        # ---------------------------------------------------------
        # PAGE 2+: DETAILED ROOM BREAKDOWNS
        # ---------------------------------------------------------
        
        doc.add_page_break()
        add_manual_heading('2. Detailed Room Specifications', 16)
        doc.add_paragraph("The following section details the specific technology and pricing for each room.")

        for room in room_list:
            name = room['name']
            data = room['config']
            
            # --- DETECT 55" DISPLAY ---
            # Check if '55' is in the display model name to trigger the specific text
            is_55_inch = "55" in data['display_model'] or "55" in data['tier_name']

            table_room = doc.add_table(rows=0, cols=5)
            table_room.style = 'Table Grid'
            
            table_room.autofit = False 
            table_room.allow_autofit = False
            
            table_room.columns[0].width = Cm(3.0)
            table_room.columns[1].width = Cm(1.3)
            table_room.columns[2].width = Cm(8.5)
            table_room.columns[3].width = Cm(2.5)
            table_room.columns[4].width = Cm(2.7)
            
            # --- ROW 1: ROOM TITLE ---
            row_hdr_obj = table_room.add_row()
            format_row(row_hdr_obj, 1.0)
            row_hdr = row_hdr_obj.cells
            
            row_hdr[0].merge(row_hdr[4])
            row_hdr[0].text = f"ROOM: {name} - {data['tier_name']} (Furthest Participant: {room['distance']}m)"
            shade_cell(row_hdr[0], "1F4E79")
            run = row_hdr[0].paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(11)

            # --- OPTIONAL: TEXT ABOVE IMAGE (Only for 55") ---
            if is_55_inch:
                row_text_top = table_room.add_row()
                # Auto height for text
                row_text_top.height_rule = WD_ROW_HEIGHT_RULE.AUTO 
                rt = row_text_top.cells
                rt[0].merge(rt[4])
                
                text_55_top = (
                    "This 4P Meeting Room shall have a 55” display mounted at 1200mm AFFL, carefully positioned to provide optimal viewing for all participants. "
                    "This height allows for easy viewing by seated participants close to the display.\n"
                    "An HDMI wall plate will be installed beneath table, with a supplied HDMI fly lead for easy device connectivity via a table box. "
                    "The display will operate on an automatic timer to turn on and off as needed, with manual control available via remote. "
                    "If a conferencing system is included, it will serve as the primary control source, streamlining operation and ensuring seamless integration of visual and audio communication."
                )
                rt[0].text = text_55_top
                rt[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                # Add a little padding visual (not real padding, just space)
                rt[0].paragraphs[0].paragraph_format.space_before = Pt(6)
                rt[0].paragraphs[0].paragraph_format.space_after = Pt(6)

            # --- ROW 2: IMAGE PLACEHOLDER ---
            row_img_obj = table_room.add_row()
            format_row(row_img_obj, 5.0)
            row_img = row_img_obj.cells
            
            row_img[0].merge(row_img[4])
            row_img[0].text = "[PASTE FLOOR PLAN IMAGE HERE]"
            row_img[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_img[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(200, 200, 200)

            # --- OPTIONAL: TEXT BELOW IMAGE (Only for 55") ---
            if is_55_inch:
                row_text_bot = table_room.add_row()
                row_text_bot.height_rule = WD_ROW_HEIGHT_RULE.AUTO
                rb = row_text_bot.cells
                rb[0].merge(rb[4])

                text_55_bot = (
                    "A Microsoft Teams Certified Cisco conferencing system mounted under the display, with a touch screen console on the conference table for intuitive operation. "
                    "It includes a Bring Your Own Device (BYOD) solution that enables users to connect personal laptops and seamlessly transfer control to their own devices.\n\n"
                    "The room is to be equipped with a Cisco Room Bar, which has a 12MP camera with a 120-degree horizontal field of view, noise cancelling microphone array, and stereo loudspeakers.\n\n"
                    "Connectivity and Power (Works in Association):\n"
                    "• Base Provision:\n"
                    "    o Behind the display, one double GPO and one data point is required.\n"
                    "    o A shielded Cat6A cable from the display to the table box for BYOD connectivity.\n"
                    "• Options:\n"
                    "    o Should Conferencing be required, two data points shall be behind the display, with a data at the table for the touch screen console, and a further data at the door for a room booking panel. "
                    "A cable path from behind the display to the ceiling shall be required for the microphone network cable which direct connects to the room bar pro.\n"
                    "    o Further cable extensions or other provisions can be supplied depending on room needs."
                )
                rb[0].text = text_55_bot
                rb[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                rb[0].paragraphs[0].paragraph_format.space_before = Pt(6)
                rb[0].paragraphs[0].paragraph_format.space_after = Pt(6)
            
            # --- ROW 3: COLUMN HEADERS ---
            row_cols_obj = table_room.add_row()
            format_row(row_cols_obj, 0.9)
            row_cols = row_cols_obj.cells
            
            row_cols[0].text = "Item"
            row_cols[1].text = "Qty"
            row_cols[2].text = "Description / Model"
            row_cols[3].text = "Unit Cost"
            row_cols[4].text = "Total"
            
            for c in row_cols:
                shade_cell(c, "D9E2F3")
                c.paragraphs[0].runs[0].bold = True
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # --- SECTION 1: CISCO ---
            row_sec1_obj = table_room.add_row()
            format_row(row_sec1_obj, 0.8)
            row_sec1 = row_sec1_obj.cells
            
            row_sec1[0].merge(row_sec1[4])
            row_sec1[0].text = "1. Data#3 Supply Scope (Cisco Hardware)"
            shade_cell(row_sec1[0], "E7E6E6")
            row_sec1[0].paragraphs[0].runs[0].bold = True

            for item in data['cisco_items']:
                if not item: continue # Skip empty strings
                r_obj = table_room.add_row()
                format_row(r_obj, 0.9)
                r = r_obj.cells
                
                r[0].text = "Video Conf"
                r[1].text = "1"
                r[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                r[2].text = item 
                r[3].text = "Excl."
                r[4].text = "Excl."
                r[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # --- SECTION 2: ALDER ---
            row_sec2_obj = table_room.add_row()
            format_row(row_sec2_obj, 0.8)
            row_sec2 = row_sec2_obj.cells
            
            row_sec2[0].merge(row_sec2[4])
            row_sec2[0].text = "2. Alder Technology Supply Scope"
            shade_cell(row_sec2[0], "E7E6E6")
            row_sec2[0].paragraphs[0].runs[0].bold = True

            def add_spec_row(item_cat, qty, desc_text, price_val):
                r_obj = table_room.add_row()
                format_row(r_obj, 0.9)
                r = r_obj.cells
                
                r[0].text = item_cat
                r[1].text = str(qty)
                r[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                r[2].text = desc_text
                r[3].text = f"${price_val:,.2f}"
                r[4].text = f"${(price_val * qty):,.2f}"
                r[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            add_spec_row("Visual Display", 1, data['display_model'], data['display_price'])
            add_spec_row("Mounting", 1, data['mount_model'], data['mount_price'])
            add_spec_row("Cabling", 1, data['cables_misc'], data['cables_price'])
            
            # Service Row
            r_svc_obj = table_room.add_row()
            format_row(r_svc_obj, 0.9)
            r_svc = r_svc_obj.cells
            
            r_svc[0].text = "Services"
            r_svc[1].text = "1"
            r_svc[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_svc[2].text = "Professional Services: Installation, Staging & PM"
            r_svc[3].text = f"${data['service_price']:,.2f}"
            r_svc[4].text = f"${data['service_price']:,.2f}"
            r_svc[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_svc[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            upfront_subtotal = data['display_price'] + data['mount_price'] + data['cables_price'] + data['service_price']
            
            row_sub_obj = table_room.add_row()
            format_row(row_sub_obj, 1.0)
            row_sub = row_sub_obj.cells
            
            row_sub[0].merge(row_sub[3])
            row_sub[0].text = "SUB-TOTAL (EX GST):"
            row_sub[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_sub[0].paragraphs[0].runs[0].bold = True
            row_sub[4].text = f"${upfront_subtotal:,.2f}"
            row_sub[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_sub[4].paragraphs[0].runs[0].bold = True

            # --- SECTION 3: MANAGED SERVICES ---
            row_sec3_obj = table_room.add_row()
            format_row(row_sec3_obj, 0.8)
            row_sec3 = row_sec3_obj.cells
            
            row_sec3[0].merge(row_sec3[4])
            row_sec3[0].text = "3. Managed Services"
            shade_cell(row_sec3[0], "E7E6E6")
            row_sec3[0].paragraphs[0].runs[0].bold = True

            r_msa_obj = table_room.add_row()
            format_row(r_msa_obj, 0.9)
            r_msa = r_msa_obj.cells
            
            r_msa[0].text = "Support"
            r_msa[1].text = "1"
            r_msa[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_msa[2].text = "Managed Service Agreement - Year 1 (Annual Billing)"
            r_msa[3].text = f"${data['ms_annual']:,.2f}"
            r_msa[4].text = f"${data['ms_annual']:,.2f}"
            r_msa[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_msa[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            doc.add_paragraph("")

        # ---------------------------------------------------------
        # FINAL SECTION
        # ---------------------------------------------------------
        doc.add_page_break()
        add_manual_heading('Further Options', 14)
        
        room_count = len(room_list)
        panel_cost = 2200.00
        total_panel_cost = room_count * panel_cost
        
        table_opt = doc.add_table(rows=2, cols=4)
        table_opt.style = 'Table Grid'
        
        # Format Headers
        hdr_opt_row = table_opt.rows[0]
        format_row(hdr_opt_row, 1.0)
        headers = hdr_opt_row.cells
        headers[0].text = "Item"
        headers[1].text = "Quantity"
        headers[2].text = "Cost per Room"
        headers[3].text = "Total"
        
        for c in headers: 
            shade_cell(c, "D9E2F3")
            c.paragraphs[0].runs[0].bold = True
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Format Data Row
        row_opt_obj = table_opt.rows[1]
        format_row(row_opt_obj, 0.9)
        row = row_opt_obj.cells
        
        row[0].text = "Room Booking Panel"
        row[1].text = str(room_count)
        row[2].text = f"${panel_cost:,.2f}"
        row[3].text = f"${total_panel_cost:,.2f}"
        
        doc.add_paragraph("\n") 

        add_manual_heading('Managed Service Agreement', 14)
        doc.add_paragraph(
            "Pricing excludes GST and is charged annually with an increase each year of 4% or CPI whichever is the greater. "
            "Acceptance of a 60 month agreement upfront locks pricing for the five year term with no increase for CPI. "
            "Pricing includes all cloud monitoring hosting charges and any onsite support required."
        )

        add_manual_heading('Exclusions', 14)
        
        exclusions_text = (
            "We exclude the following items from our installation. "
            "We exclude all power and data, plus building works. All works required shall be identified and must be completed prior to our installers attending site. "
            "Any Cat6/6A Tie Lines between Audio Visual Locations (e.g. behind display to table box/floor box) shall be by the data contractor. "
            "We exclude all decommissioning, disposal of equipment and make good works. "
            "We exclude all height access equipment, and all furniture protection equipment/coverings. "
            "All Webex/Teams/Exchange credentials must be provided prior to attending site. Admin credentials for any Webex/Teams endpoint must be provided to Alder Technology for the duration of any deployment including temporary administrator cloud tenant access and other software admin access for the duration of deployment. "
            "The network must be active and configured for Teams/Webex prior to attending site. "
            "All external services provided by the client or their nominated systems integrator, including Microsoft Teams Room/Webex accounts, Azure Intune registrations and specific deployment requirements, Exchange, Skype for Business and Microsoft security and compliance requirements and Fast track or Peering ISP plans are the responsibility of the client. "
            "Software updates and the impact on the hardware, user experience and the operational state of the system are the sole responsibility of the client. Any such updates that require Alder Technology site attendance incur a Service call out fee and hourly rate at agreed hourly rates, unless a service level agreement covering these works is in place. "
            "A site planner shall be provided by Alder Technology and must be completed by the client prior to our installers attending site. "
            "Should works not be able to commence due to the above or client led delays, a call out fee may be applied. "
            "Pricing is based on a consecutive work package. Significant delays or pauses in works due to client delays or room unavailability may result in additional charges."
        )
        
        doc.add_paragraph(exclusions_text)

        doc.add_paragraph("\nThank you for your consideration. Please call me if you have any further queries.")
        doc.add_paragraph("Regards,")
        sig = doc.add_paragraph("George Coles")
        sig.runs[0].bold = True

        desktop = os.path.expanduser("~/Desktop")
        save_folder = os.path.join(desktop, "Alder_Quotes")
        if not os.path.exists(save_folder):
            os.makedirs(save_folder)

        safe_client_name = "".join([c for c in client_name if c.isalpha() or c.isdigit() or c==' ']).rstrip()
        timestamp = datetime.now().strftime("%H-%M-%S")
        filename = f"Alder_Quote_{safe_client_name}_{timestamp}.docx"
        full_path = os.path.join(save_folder, filename)
        
        doc.save(full_path)
        return full_path

    # ==========================================
    # PART 2: THE USER INTERFACE (GUI)
    # ==========================================

    ctk.set_appearance_mode("Light") 
    ctk.set_default_color_theme("green") 

    app = ctk.CTk()
    app.geometry("500x750")
    app.title("Alder Technology - Quoting Tool")
    app.configure(fg_color="#E5E5E5")

    # --- INITIALIZE DATA ---
    PRICELIST_DATA, EXCEL_ERROR = load_pricelist_from_excel()

    # --- LOGO HANDLING ---
    script_dir = os.path.dirname(os.path.abspath(__file__))
    logo_path = os.path.join(script_dir, "alder_logo.png")
    
    logo_image = None
    if os.path.exists(logo_path):
        try:
            pil_image = Image.open(logo_path)
            logo_image = ctk.CTkImage(light_image=pil_image, dark_image=pil_image, size=(250, 180))
        except Exception as e:
            print(f"Error loading logo: {e}")

    def on_generate_click():
        lbl_status.configure(text="Processing...", text_color="black")
        app.update()
        
        if EXCEL_ERROR:
            lbl_status.configure(text=f"EXCEL ERROR: {EXCEL_ERROR}", text_color="#FF5555")
            return

        client = entry_client.get()
        if not client:
            lbl_status.configure(text="Error: Enter Client Name", text_color="#FF5555")
            return

        raw_text = txt_rooms.get("0.0", "end")
        lines = raw_text.split('\n')
        valid_rooms = []
        
        try:
            for line in lines:
                line = line.strip()
                if not line: continue 
                if "\t" in line: line = line.replace("\t", ",") 
                if "," not in line: continue 
                    
                parts = line.split(',')
                r_name = parts[0].strip()
                r_dist_str = parts[-1].lower().replace('m', '').strip()
                
                try:
                    r_dist = float(r_dist_str)
                except ValueError:
                    continue 
                
                # --- NEW: PASS PRICELIST TO CONFIG FUNCTION ---
                config = get_room_configuration(r_dist, PRICELIST_DATA)
                
                if config:
                    valid_rooms.append({'name': r_name, 'distance': r_dist, 'config': config})
                else:
                    print(f"Skipping room {r_name} - distance {r_dist}m exceeds all tiers in Excel.")
            
            if len(valid_rooms) == 0:
                lbl_status.configure(text="Error: No valid rooms found within range.", text_color="#FF5555")
                return

            filepath = generate_multi_room_proposal(client, valid_rooms)
            lbl_status.configure(text=f"SUCCESS!\nSaved to Desktop/Alder_Quotes:\n{os.path.basename(filepath)}", text_color="#009A44")
            
            try: os.startfile(os.path.dirname(filepath))
            except: pass
            
        except PermissionError:
            lbl_status.configure(text="ERROR: PERMISSION DENIED.\nClose Word and try again.", text_color="#FF5555")
        except Exception as e:
            lbl_status.configure(text=f"Error: {str(e)}", text_color="#FF5555")

    # --- GUI LAYOUT ---
    
    main_frame = ctk.CTkFrame(app, corner_radius=15, fg_color="white")
    main_frame.pack(pady=20, padx=20, fill="both", expand=True)

    if logo_image:
        lbl_logo = ctk.CTkLabel(main_frame, image=logo_image, text="")
        lbl_logo.pack(pady=(20, 5))
    else:
        lbl_title = ctk.CTkLabel(main_frame, text="ALDER TECHNOLOGY", font=("Arial Black", 24), text_color="#009A44")
        lbl_title.pack(pady=(20, 5))

    lbl_subtitle = ctk.CTkLabel(main_frame, text="Multi-Room Proposal Generator", font=("Roboto", 13), text_color="gray")
    lbl_subtitle.pack(pady=(0, 20))

    # Check Excel Status for GUI
    if EXCEL_ERROR:
        lbl_db_status = ctk.CTkLabel(main_frame, text=f"⚠️ {EXCEL_ERROR}", text_color="red", font=("Arial", 11, "bold"))
    else:
        lbl_db_status = ctk.CTkLabel(main_frame, text=f"✔ Loaded {len(PRICELIST_DATA)} tiers from Excel", text_color="green", font=("Arial", 11))
    lbl_db_status.pack(pady=(0, 10))

    entry_client = ctk.CTkEntry(main_frame, placeholder_text="Client Name (e.g. Acme Corp)", width=350, height=40)
    entry_client.pack(pady=10)

    lbl_instr = ctk.CTkLabel(main_frame, text="Paste Excel Data below (Room Name | Furthest Participant)", text_color="gray", font=("Roboto", 12))
    lbl_instr.pack(pady=(10, 2))

    txt_rooms = ctk.CTkTextbox(main_frame, width=350, height=200, corner_radius=10, border_width=1, border_color="#CCC")
    txt_rooms.pack(pady=5)
    
    btn_generate = ctk.CTkButton(
        main_frame, 
        text="GENERATE PROPOSAL", 
        command=on_generate_click, 
        width=350, 
        height=50, 
        font=("Roboto", 15, "bold"),
        fg_color="#009A44", 
        hover_color="#007a36",
        text_color="white",
        state="disabled" if EXCEL_ERROR else "normal"
    )
    btn_generate.pack(pady=20)

    lbl_status = ctk.CTkLabel(main_frame, text="Ready", text_color="black", font=("Roboto", 12))
    lbl_status.pack(pady=10)
    
    lbl_footer = ctk.CTkLabel(app, text="v3.0 (Excel Linked) | Alder Technology Internal", text_color="#666", font=("Arial", 10))
    lbl_footer.pack(side="bottom", pady=10)

    app.mainloop()

# --- CRASH CATCHER ---
except Exception as e:
    root = tkinter.Tk()
    root.withdraw()
    error_msg = str(traceback.format_exc())
    tkinter.messagebox.showerror("Program Error", f"The program crashed:\n\n{error_msg}")
