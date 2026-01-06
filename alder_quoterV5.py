import traceback
import tkinter
import tkinter.messagebox

# --- CRASH REPORTER WRAPPER ---
try:
    import customtkinter as ctk
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from datetime import datetime
    import os
    import sys

    # ==========================================
    # PART 1: THE LOGIC & PRICING DATABASE
    # ==========================================

    def get_room_configuration(distance):
        """
        Returns the specific Bill of Materials and Pricing based on room depth.
        """
        # --- TIER 1: UP TO 3.0M ---
        if distance <= 3.0:
            return {
                "tier_name": "Small Meeting Space",
                # CISCO (Data3 Supply)
                "cisco_items": [
                    "Cisco Room Bar",
                    "Integrated Camera",
                    "Integrated Audio"
                ],
                # ALDER HARDWARE
                "display_model": "LG 55UL3J-B (55\" 4K UHD, 400nits, WebOS)",
                "display_price": 1100.00,
                "mount_model": "Venturi VP-F80 (Suit 50\"-75\")",
                "mount_price": 65.00,
                "cables_misc": " HDMI & Patch Leads", # Added minor consumables
                "cables_price": 50.00,
                # ALDER SERVICES
                "service_price": 3000.00,
                "ms_annual": 1200.00
            }
        # --- TIER 2: UP TO 4.5M ---
        elif distance <= 4.5:
            return {
                "tier_name": "Medium Meeting Space",
                "cisco_items": [
                    "Cisco Room Bar",
                    "Integrated Camera",
                    "1x Cisco Table Microphone Pro"
                ],
                "display_model": "LG 65UL3J-B (65\" 4K UHD, 400nits, WebOS)",
                "display_price": 1500.00,
                "mount_model": "Venturi VP-F80 (Suit 50\"-75\")",
                "mount_price": 65.00,
                "cables_misc": "HDMI & Patch Leads",
                "cables_price": 50.00,
                "service_price": 3000.00,
                "ms_annual": 1200.00
            }
        # --- TIER 3: UP TO 5.5M ---
        elif distance <= 5.5:
            return {
                "tier_name": "Large Meeting Space",
                "cisco_items": [
                    "Cisco Room Bar Pro",
                    "Integrated Dual Camera",
                    "1x Cisco Ceiling Microphone Pro"
                ],
                "display_model": "LG 75UL3J-B (75\" 4K UHD, 400nits, WebOS)",
                "display_price": 2200.00,
                "mount_model": "Venturi VP-F80 (Suit 50\"-75\")",
                "mount_price": 65.00,
                "cables_misc": "HDMI, Patch Leads & Mount Fixings",
                "cables_price": 80.00,
                "service_price": 3000.00,
                "ms_annual": 1500.00
            }
        # --- TIER 4: UP TO 6.5M ---
        elif distance <= 6.5:
            return {
                "tier_name": "Extra Large Space",
                "cisco_items": [
                    "Cisco Room Bar Pro",
                    "Integrated Dual Camera",
                    "1x Cisco Ceiling Microphone Pro"
                ],
                "display_model": "LG 86UL3J-B (86\" 4K UHD, 330nits, WebOS)",
                "display_price": 3300.00,
                "mount_model": "VP-F100 (Suit 82\"-98\")",
                "mount_price": 100.00,
                "cables_misc": "HDMI, Patch Leads & Heavy Duty Fixings",
                "cables_price": 100.00,
                "service_price": 3500.00,
                "ms_annual": 1500.00
            }
        # --- TIER 5: UP TO 7.5M ---
        elif distance <= 7.5:
            return {
                "tier_name": "Boardroom",
                "cisco_items": [
                    "Cisco Kit EQ + AV Integrator License",
                    "Cisco Quad Cam",
                    "2x Cisco Ceiling Mic Pro",
                    "6-8x Shure Ceiling Speakers"
                ],
                "display_model": "LG 98UM5K (98\" 4K UHD, 500nits)",
                "display_price": 7500.00,
                "mount_model": "VP-F100 (Suit 82\"-98\")",
                "mount_price": 100.00,
                "cables_misc": "HDMI, Patch Leads, Speaker Cable, Fixings",
                "cables_price": 200.00,
                "service_price": 4000.00,
                "ms_annual": 3000.00
            }
        else:
            return None

    def generate_multi_room_proposal(client_name, room_list):
        doc = Document()
        
        # Style margins
        section = doc.sections[0]
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

        # ---------------------------------------------------------
        # PAGE 1: EXECUTIVE SUMMARY & MASTER SCHEDULE
        # ---------------------------------------------------------
        
        # Header
        header = doc.add_heading(f'AV Proposal: {client_name}', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f'Prepared by: Alder Technology')
        doc.add_paragraph(f'Date: {datetime.now().strftime("%d/%m/%Y")}')
        doc.add_paragraph(f'Total Rooms Scoped: {len(room_list)}')
        doc.add_paragraph('------------------------------------------------------')

        # Intro
        doc.add_heading('Partnership Overview', level=2)
        doc.add_paragraph(
            "Alder Technology is pleased to partner with Data#3 to provide this solution. "
            "This document is split into two sections:\n"
            "1. A Master Financial Schedule summarizing all rooms.\n"
            "2. Detailed Bill of Materials for each specific room.\n\n"
            "Please note: Cisco hardware is listed for engineering reference but is to be supplied and priced by Data#3."
        )

        # Master Table
        doc.add_heading('1. Master Room Schedule & Pricing', level=2)
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Headers
        hdr = table.rows[0].cells
        hdr[0].text = "Room Name"
        hdr[1].text = "Classification"
        hdr[2].text = "Alder Hardware"
        hdr[3].text = "Services + MS (5Yr)"
        hdr[4].text = "Room Total (Ex GST)"

        grand_total_project = 0

        # LOOP FOR MASTER TABLE
        for room in room_list:
            name = room['name']
            dist = room['distance']
            data = room['config']
            
            # Calculations
            hw_cost = data['display_price'] + data['mount_price'] + data['cables_price']
            svc_cost = data['service_price']
            ms_5yr = data['ms_annual'] * 5
            room_total = hw_cost + svc_cost + ms_5yr
            grand_total_project += room_total

            # Add Row
            row = table.add_row().cells
            row[0].text = name
            row[1].text = f"{data['tier_name']}\n({dist}m)"
            row[2].text = f"${hw_cost:,.0f}"
            row[3].text = f"${(svc_cost + ms_5yr):,.0f}"
            row[4].text = f"${room_total:,.2f}"

        # Grand Total Text
        doc.add_paragraph("\n")
        p_total = doc.add_paragraph()
        p_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        runner = p_total.add_run(f"GRAND TOTAL PROJECT VALUE (EX GST): ${grand_total_project:,.2f}")
        runner.bold = True
        runner.font.size = Pt(16)
        runner.font.color.rgb = RGBColor(0, 102, 204)
        
        # ---------------------------------------------------------
        # PAGE 2+: DETAILED ROOM BREAKDOWNS
        # ---------------------------------------------------------
        
        doc.add_page_break()
        doc.add_heading('2. Detailed Room Specifications', level=1)
        doc.add_paragraph("The following pages detail the specific technology and pricing for each room.")

        for room in room_list:
            # Data Setup
            name = room['name']
            data = room['config']
            
            # Visual separator between rooms
            doc.add_paragraph("\n------------------------------------------------------")
            doc.add_heading(f"ROOM: {name} ({room['distance']}m Depth)", level=2)
            
            # --- SUB-TABLE 1: CISCO (DATA3) ---
            doc.add_heading('Data#3 Supply Scope (Cisco Systems)', level=3)
            table_c = doc.add_table(rows=1, cols=2)
            table_c.style = 'Light Shading Accent 1' # Grey table to signify "Not our price"
            table_c.rows[0].cells[0].text = "Cisco Component"
            table_c.rows[0].cells[1].text = "Status"
            
            for item in data['cisco_items']:
                row = table_c.add_row().cells
                row[0].text = item
                row[1].text = "Supplied by Data#3"

            # --- SUB-TABLE 2: ALDER HARDWARE ---
            doc.add_paragraph("") # Spacing
            doc.add_heading('Alder Technology Supply Scope', level=3)
            table_a = doc.add_table(rows=1, cols=3)
            table_a.style = 'Table Grid'
            
            # Headers
            h = table_a.rows[0].cells
            h[0].text = "Item Category"
            h[1].text = "Description / Model"
            h[2].text = "Unit Price"

            # 1. Display
            r = table_a.add_row().cells
            r[0].text = "Visual Display"
            r[1].text = data['display_model']
            r[2].text = f"${data['display_price']:,.2f}"

            # 2. Mount
            r = table_a.add_row().cells
            r[0].text = "Mounting Hardware"
            r[1].text = data['mount_model']
            r[2].text = f"${data['mount_price']:,.2f}"

            # 3. Cables
            r = table_a.add_row().cells
            r[0].text = "Cabling & Consumables"
            r[1].text = data['cables_misc']
            r[2].text = f"${data['cables_price']:,.2f}"

            # 4. Services (Breakdown)
            # We list the services as one line item price, but describe inclusions
            r = table_a.add_row().cells
            r[0].text = "Professional Services"
            r[1].text = "Includes: Offsite Staging, Installation, Room Coordination, Commissioning, PM & Documentation."
            r[2].text = f"${data['service_price']:,.2f}"

            # 5. Managed Services
            r = table_a.add_row().cells
            r[0].text = "Managed Service (5-Year)"
            r[1].text = f"5-Year Agreement @ ${data['ms_annual']} per annum"
            r[2].text = f"${(data['ms_annual']*5):,.2f}"

            # Room Total
            this_room_total = data['display_price'] + data['mount_price'] + data['cables_price'] + data['service_price'] + (data['ms_annual']*5)
            
            # Total Row
            r = table_a.add_row().cells
            r[0].text = ""
            r[1].text = "TOTAL FOR THIS ROOM (EX GST):"
            r[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r[2].text = f"${this_room_total:,.2f}"
            r[2].paragraphs[0].runs[0].bold = True

            doc.add_paragraph("") # Space after room

        # --- SAVE LOGIC ---
        desktop = os.path.expanduser("~/Desktop")
        save_folder = os.path.join(desktop, "Alder_Quotes")
        if not os.path.exists(save_folder):
            os.makedirs(save_folder)

        safe_client_name = "".join([c for c in client_name if c.isalpha() or c.isdigit() or c==' ']).rstrip()
        timestamp = datetime.now().strftime("%H-%M-%S")
        filename = f"Alder_Quote_{safe_client_name}_{timestamp}.docx"
        full_path = os.path.join(save_folder, filename)
        
        doc.save(full_path)
        return full_path

    # ==========================================
    # PART 2: THE USER INTERFACE (GUI)
    # ==========================================

    ctk.set_appearance_mode("Dark")
    ctk.set_default_color_theme("blue")

    app = ctk.CTk()
    app.geometry("600x700")
    app.title("Alder Technology - Multi-Room Quoter")

    def on_generate_click():
        lbl_status.configure(text="Processing...", text_color="white")
        
        client = entry_client.get()
        if not client:
            lbl_status.configure(text="Error: Enter Client Name", text_color="#FF5555")
            return

        raw_text = txt_rooms.get("0.0", "end")
        lines = raw_text.split('\n')
        valid_rooms = []
        
        try:
            for line in lines:
                line = line.strip()
                if not line: continue 
                if "\t" in line: line = line.replace("\t", ",") # Excel fix
                if "," not in line: continue 
                    
                parts = line.split(',')
                r_name = parts[0].strip()
                r_dist_str = parts[-1].lower().replace('m', '').strip()
                
                try:
                    r_dist = float(r_dist_str)
                except ValueError:
                    continue 
                
                config = get_room_configuration(r_dist)
                if config:
                    valid_rooms.append({'name': r_name, 'distance': r_dist, 'config': config})
            
            if len(valid_rooms) == 0:
                lbl_status.configure(text="Error: No valid rooms found.\nFormat: 'Name, Distance'", text_color="#FF5555")
                return

            filepath = generate_multi_room_proposal(client, valid_rooms)
            lbl_status.configure(text=f"SUCCESS!\nSaved to Desktop/Alder_Quotes:\n{os.path.basename(filepath)}", text_color="#00FF00")
            
            try: os.startfile(os.path.dirname(filepath))
            except: pass
            
        except PermissionError:
            lbl_status.configure(text="ERROR: PERMISSION DENIED.\nClose Word and try again.", text_color="#FF5555")
        except Exception as e:
            lbl_status.configure(text=f"Error: {str(e)}", text_color="#FF5555")

    # --- UI ELEMENTS ---
    lbl_title = ctk.CTkLabel(app, text="Alder Multi-Room Quoter", font=("Roboto Medium", 24))
    lbl_title.pack(pady=15)

    entry_client = ctk.CTkEntry(app, placeholder_text="Client Name (e.g. Acme Corp)", width=400)
    entry_client.pack(pady=5)

    lbl_instr = ctk.CTkLabel(app, text="Paste Excel Data below (Name | Depth)", text_color="gray")
    lbl_instr.pack(pady=(15, 5))

    txt_rooms = ctk.CTkTextbox(app, width=400, height=300)
    txt_rooms.pack(pady=5)

    btn_generate = ctk.CTkButton(app, text="GENERATE MASTER PROPOSAL", command=on_generate_click, width=400, height=50, font=("Roboto Medium", 14))
    btn_generate.pack(pady=20)

    lbl_status = ctk.CTkLabel(app, text="Ready...", font=("Roboto", 12))
    lbl_status.pack(pady=10)

    app.mainloop()

# --- CRASH CATCHER ---
except Exception as e:
    root = tkinter.Tk()
    root.withdraw()
    error_msg = str(traceback.format_exc())
    tkinter.messagebox.showerror("Program Error", f"The program crashed:\n\n{error_msg}")
