import traceback
import tkinter
import tkinter.messagebox

# --- CRASH REPORTER WRAPPER ---
try:
    import customtkinter as ctk
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from datetime import datetime
    import os
    import sys
    from PIL import Image

    # ==========================================
    # PART 1: THE LOGIC & PRICING DATABASE
    # ==========================================

    def get_room_configuration(distance):
        # --- TIER 1: UP TO 3.0M ---
        if distance <= 3.0:
            return {
                "tier_name": "Small Meeting Space",
                "cisco_items": ["Cisco Room Bar", "Integrated Camera", "Integrated Audio"],
                "display_model": "LG 55UL3J-B (55\" 4K UHD, 400nits, WebOS)",
                "display_price": 1100.00,
                "mount_model": "Venturi VP-F80 (Suit 50\"-75\")",
                "mount_price": 65.00,
                "cables_misc": "HDMI & Patch Leads", 
                "cables_price": 50.00,
                "service_price": 3000.00,
                "ms_annual": 1200.00
            }
        # --- TIER 2: UP TO 4.5M ---
        elif distance <= 4.5:
            return {
                "tier_name": "Medium Meeting Space",
                "cisco_items": ["Cisco Room Bar", "Integrated Camera", "1x Cisco Table Microphone Pro"],
                "display_model": "LG 65UL3J-B (65\" 4K UHD, 400nits, WebOS)",
                "display_price": 1500.00,
                "mount_model": "Venturi VP-F80 (Suit 50\"-75\")",
                "mount_price": 65.00,
                "cables_misc": "HDMI & Patch Leads",
                "cables_price": 50.00,
                "service_price": 3000.00,
                "ms_annual": 1200.00
            }
        # --- TIER 3: UP TO 5.5M ---
        elif distance <= 5.5:
            return {
                "tier_name": "Large Meeting Space",
                "cisco_items": ["Cisco Room Bar Pro", "Integrated Dual Camera", "1x Cisco Ceiling Microphone Pro"],
                "display_model": "LG 75UL3J-B (75\" 4K UHD, 400nits, WebOS)",
                "display_price": 2200.00,
                "mount_model": "Venturi VP-F80 (Suit 50\"-75\")",
                "mount_price": 65.00,
                "cables_misc": "HDMI, Patch Leads & Mount Fixings",
                "cables_price": 80.00,
                "service_price": 3000.00,
                "ms_annual": 1500.00
            }
        # --- TIER 4: UP TO 6.5M ---
        elif distance <= 6.5:
            return {
                "tier_name": "Extra Large Space",
                "cisco_items": ["Cisco Room Bar Pro", "Integrated Dual Camera", "1x Cisco Ceiling Microphone Pro"],
                "display_model": "LG 86UL3J-B (86\" 4K UHD, 330nits, WebOS)",
                "display_price": 3300.00,
                "mount_model": "VP-F100 (Suit 82\"-98\")",
                "mount_price": 100.00,
                "cables_misc": "HDMI, Patch Leads & Heavy Duty Fixings",
                "cables_price": 100.00,
                "service_price": 3500.00,
                "ms_annual": 1500.00
            }
        # --- TIER 5: BOARDROOM (Fixed Range to catch 7.8m) ---
        elif distance <= 15.0:
            return {
                "tier_name": "Boardroom",
                "cisco_items": ["Cisco Kit EQ + AV Integrator License", "Cisco Quad Cam", "2x Cisco Ceiling Mic Pro", "6-8x Shure Ceiling Speakers"],
                "display_model": "LG 98UM5K (98\" 4K UHD, 500nits)",
                "display_price": 7500.00,
                "mount_model": "VP-F100 (Suit 82\"-98\")",
                "mount_price": 100.00,
                "cables_misc": "HDMI, Patch Leads, Speaker Cable, Fixings",
                "cables_price": 200.00,
                "service_price": 4000.00,
                "ms_annual": 3000.00
            }
        else:
            return None

    def generate_multi_room_proposal(client_name, room_list):
        
        # --- TEMPLATE LOADING ---
        script_dir = os.path.dirname(os.path.abspath(__file__))
        template_path = os.path.join(script_dir, "Alder_Template.docx")
        
        if os.path.exists(template_path):
            doc = Document(template_path)
            doc.add_page_break() 
        else:
            doc = Document()

        # --- SET FONT TO HELVETICA ---
        try:
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Helvetica'
            font.size = Pt(10)
        except:
            pass 

        # Style margins
        try:
            section = doc.sections[0]
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
        except: pass

        # --- HELPER: SHADE TABLE CELLS ---
        def shade_cell(cell, color_hex):
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
            cell._tc.get_or_add_tcPr().append(shading_elm)

        # --- HELPER: MANUAL HEADING ---
        def add_manual_heading(text, size, color_rgb=None):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Helvetica'
            run.font.size = Pt(size)
            if color_rgb:
                run.font.color.rgb = color_rgb
            return p

        # ---------------------------------------------------------
        # PAGE 1: EXECUTIVE SUMMARY & MASTER LIST
        # ---------------------------------------------------------
        
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_t = p_title.add_run(f'AV Proposal: {client_name}')
        run_t.bold = True
        run_t.font.size = Pt(26)
        
        doc.add_paragraph(f'Prepared by: Alder Technology')
        doc.add_paragraph(f'Date: {datetime.now().strftime("%d/%m/%Y")}')
        doc.add_paragraph(f'Total Rooms Scoped: {len(room_list)}')
        doc.add_paragraph('------------------------------------------------------')

        add_manual_heading('Partnership Overview', 14)
        doc.add_paragraph(
            "Alder Technology is pleased to partner with Data#3 to provide this solution. "
            "This document is split into two sections:\n"
            "1. A Master Financial Summary (Hardware + Year 1 Services).\n"
            "2. Detailed Bill of Materials for each specific room.\n\n"
            "Please note: Cisco hardware is listed for engineering reference but is to be supplied and priced by Data#3."
        )

        add_manual_heading('1. Master Room Summary & Pricing', 14)
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        hdr = table.rows[0].cells
        hdr[0].text = "Room Name"
        hdr[1].text = "Classification"
        hdr[2].text = "Supply & Services" 
        hdr[3].text = "Managed Service (Year 1)"
        hdr[4].text = "Total Year 1 (Ex GST)"

        for cell in hdr:
            shade_cell(cell, "D9E2F3") # Light Blue
            cell.paragraphs[0].runs[0].bold = True

        grand_total_project = 0

        for room in room_list:
            name = room['name']
            dist = room['distance']
            data = room['config']
            
            upfront_cost = data['display_price'] + data['mount_price'] + data['cables_price'] + data['service_price']
            ms_annual = data['ms_annual'] 
            
            room_total_y1 = upfront_cost + ms_annual
            grand_total_project += room_total_y1

            row = table.add_row().cells
            row[0].text = name
            row[1].text = f"{data['tier_name']} ({dist}m)"
            row[2].text = f"${upfront_cost:,.0f}"
            row[3].text = f"${ms_annual:,.0f}"
            row[4].text = f"${room_total_y1:,.2f}"
            
            for i in range(2, 5):
                row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph("\n")
        p_total = doc.add_paragraph()
        p_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        runner = p_total.add_run(f"TOTAL YEAR 1 PROJECT VALUE (EX GST): ${grand_total_project:,.2f}")
        runner.bold = True
        runner.font.size = Pt(16)
        runner.font.color.rgb = RGBColor(0, 102, 204)
        
        doc.add_paragraph("(Includes Alder Hardware, Installation Services, and 1st Year Managed Service)")
        
        # ---------------------------------------------------------
        # PAGE 2+: DETAILED ROOM BREAKDOWNS
        # ---------------------------------------------------------
        
        doc.add_page_break()
        add_manual_heading('2. Detailed Room Specifications', 16)
        doc.add_paragraph("The following section details the specific technology and pricing for each room.")

        for room in room_list:
            name = room['name']
            data = room['config']
            
            # 5 Columns: Item | Qty | Description | Unit Cost | Total
            table_room = doc.add_table(rows=0, cols=5)
            table_room.style = 'Table Grid'
            table_room.autofit = True
            
            # --- ROW 1: ROOM TITLE ---
            row_hdr = table_room.add_row().cells
            row_hdr[0].merge(row_hdr[4])
            row_hdr[0].text = f"ROOM: {name} ({room['distance']}m) - {data['tier_name']}"
            shade_cell(row_hdr[0], "1F4E79") # Dark Blue
            run = row_hdr[0].paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(11)

            # --- ROW 2: IMAGE PLACEHOLDER ---
            row_img = table_room.add_row().cells
            row_img[0].merge(row_img[4])
            row_img[0].text = "\n\n[PASTE FLOOR PLAN IMAGE HERE]\n\n"
            row_img[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_img[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(200, 200, 200)
            
            # --- ROW 3: COLUMN HEADERS ---
            row_cols = table_room.add_row().cells
            row_cols[0].text = "Item"
            row_cols[1].text = "Qty"
            row_cols[2].text = "Description / Model" # Merged Column
            row_cols[3].text = "Unit Cost"
            row_cols[4].text = "Total"
            
            for c in row_cols:
                shade_cell(c, "D9E2F3")
                c.paragraphs[0].runs[0].bold = True
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # --- SECTION 1: CISCO ---
            row_sec1 = table_room.add_row().cells
            row_sec1[0].merge(row_sec1[4])
            row_sec1[0].text = "1. Data#3 Supply Scope (Cisco Hardware)"
            shade_cell(row_sec1[0], "E7E6E6")
            row_sec1[0].paragraphs[0].runs[0].bold = True

            for item in data['cisco_items']:
                r = table_room.add_row().cells
                r[0].text = "Video Conf"
                r[1].text = "1"
                r[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                r[2].text = f"{item} (Supplied by Data#3)"
                r[3].text = "Excl."
                r[4].text = "Excl."
                r[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # --- SECTION 2: ALDER ---
            row_sec2 = table_room.add_row().cells
            row_sec2[0].merge(row_sec2[4])
            row_sec2[0].text = "2. Alder Technology Supply Scope"
            shade_cell(row_sec2[0], "E7E6E6")
            row_sec2[0].paragraphs[0].runs[0].bold = True

            def add_spec_row(item_cat, qty, desc_text, price_val):
                r = table_room.add_row().cells
                r[0].text = item_cat
                r[1].text = str(qty)
                r[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                r[2].text = desc_text
                r[3].text = f"${price_val:,.2f}"
                r[4].text = f"${(price_val * qty):,.2f}"
                r[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            add_spec_row("Visual Display", 1, data['display_model'], data['display_price'])
            add_spec_row("Mounting", 1, data['mount_model'], data['mount_price'])
            add_spec_row("Cabling", 1, data['cables_misc'], data['cables_price'])
            
            # Service Row
            r_svc = table_room.add_row().cells
            r_svc[0].text = "Services"
            r_svc[1].text = "1"
            r_svc[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_svc[2].text = "Professional Services: Installation, Staging & PM"
            r_svc[3].text = f"${data['service_price']:,.2f}"
            r_svc[4].text = f"${data['service_price']:,.2f}"
            r_svc[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_svc[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            upfront_subtotal = data['display_price'] + data['mount_price'] + data['cables_price'] + data['service_price']
            row_sub = table_room.add_row().cells
            row_sub[0].merge(row_sub[3]) # Merge Item, Qty, Desc, Unit Cost
            row_sub[0].text = "SUB-TOTAL (EX GST):"
            row_sub[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_sub[0].paragraphs[0].runs[0].bold = True
            row_sub[4].text = f"${upfront_subtotal:,.2f}"
            row_sub[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_sub[4].paragraphs[0].runs[0].bold = True

            # --- SECTION 3: MANAGED SERVICES ---
            row_sec3 = table_room.add_row().cells
            row_sec3[0].merge(row_sec3[4])
            row_sec3[0].text = "3. Managed Services"
            shade_cell(row_sec3[0], "E7E6E6")
            row_sec3[0].paragraphs[0].runs[0].bold = True

            r_msa = table_room.add_row().cells
            r_msa[0].text = "Support"
            r_msa[1].text = "1"
            r_msa[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_msa[2].text = "Managed Service Agreement - Year 1 (Annual Billing)"
            r_msa[3].text = f"${data['ms_annual']:,.2f}"
            r_msa[4].text = f"${data['ms_annual']:,.2f}"
            r_msa[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_msa[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            doc.add_paragraph("")

        # ---------------------------------------------------------
        # FINAL SECTION
        # ---------------------------------------------------------
        doc.add_page_break()
        add_manual_heading('Further Options', 14)
        
        room_count = len(room_list)
        panel_cost = 2200.00
        total_panel_cost = room_count * panel_cost
        
        table_opt = doc.add_table(rows=2, cols=4)
        table_opt.style = 'Table Grid'
        
        headers = table_opt.rows[0].cells
        headers[0].text = "Item"
        headers[1].text = "Quantity"
        headers[2].text = "Cost per Room"
        headers[3].text = "Total"
        
        for c in headers: 
            shade_cell(c, "D9E2F3")
            c.paragraphs[0].runs[0].bold = True
        
        row = table_opt.rows[1].cells
        row[0].text = "Room Booking Panel"
        row[1].text = str(room_count)
        row[2].text = f"${panel_cost:,.2f}"
        row[3].text = f"${total_panel_cost:,.2f}"
        doc.add_paragraph("\n") 

        add_manual_heading('Managed Service Agreement', 14)
        doc.add_paragraph(
            "An example costing is provided below. Please note the below does not constitute a quote until the scope of the Managed Service Agreement is finalised. "
            "Pricing excludes GST and is charged annually with an increase each year of 4% or CPI whichever is the greater. "
            "Acceptance of a 60 month agreement upfront locks pricing for the five year term with no increase for CPI. "
            "Pricing includes all cloud monitoring hosting charges and any onsite support required."
        )

        add_manual_heading('Exclusions', 14)
        exclusions = [
            "We exclude all power and data, plus building works. All works required shall be identified and must be completed prior to our installers attending site.",
            "We exclude all height access equipment, and all furniture protection equipment/coverings.",
            "All Teams/Exchange credentials must be provided prior to attending site...",
            "The network must be active and configured for Teams prior to attending site.",
            "All external services provided by the client are the responsibility of the client.",
            "Microsoft updates and the impact on the hardware are the sole responsibility of the client.",
            "A site planner shall be provided by Alder Technology and must be completed by the client prior to our installers attending site.",
            "Should works not be able to commence due to the above or client led delays, a call out fee may be applied.",
            "Should a managed service not be engaged, a DLP period of three (3) months is applicable.",
            "Alder Technology works with partners to deliver the best product possible."
        ]
        
        for item in exclusions:
            p = doc.add_paragraph(item)
            p.paragraph_format.space_after = Pt(6)

        doc.add_paragraph("\nThank you for your consideration. Please call me if you have any further queries.")
        doc.add_paragraph("Regards,")
        sig = doc.add_paragraph("George Coles")
        sig.runs[0].bold = True

        desktop = os.path.expanduser("~/Desktop")
        save_folder = os.path.join(desktop, "Alder_Quotes")
        if not os.path.exists(save_folder):
            os.makedirs(save_folder)

        safe_client_name = "".join([c for c in client_name if c.isalpha() or c.isdigit() or c==' ']).rstrip()
        timestamp = datetime.now().strftime("%H-%M-%S")
        filename = f"Alder_Quote_{safe_client_name}_{timestamp}.docx"
        full_path = os.path.join(save_folder, filename)
        
        doc.save(full_path)
        return full_path

    # ==========================================
    # PART 2: THE NEW USER INTERFACE (GUI)
    # ==========================================

    # 1. Update Theme to Light and set Green Accent
    ctk.set_appearance_mode("Light") 
    ctk.set_default_color_theme("green") 

    app = ctk.CTk()
    app.geometry("500x750")
    app.title("Alder Technology - Quoting Tool")
    
    # 2. Set Light Grey Background for Window
    app.configure(fg_color="#E5E5E5")

    # --- LOGO HANDLING ---
    script_dir = os.path.dirname(os.path.abspath(__file__))
    logo_path = os.path.join(script_dir, "alder_logo.png")
    
    logo_image = None
    if os.path.exists(logo_path):
        try:
            pil_image = Image.open(logo_path)
            # Create CTkImage with size matching aspect ratio (approx 3:2 from image)
            logo_image = ctk.CTkImage(light_image=pil_image, dark_image=pil_image, size=(250, 180))
        except Exception as e:
            print(f"Error loading logo: {e}")

    def on_generate_click():
        lbl_status.configure(text="Processing...", text_color="black")
        app.update() # Force UI update
        
        client = entry_client.get()
        if not client:
            lbl_status.configure(text="Error: Enter Client Name", text_color="#FF5555")
            return

        raw_text = txt_rooms.get("0.0", "end")
        lines = raw_text.split('\n')
        valid_rooms = []
        
        try:
            for line in lines:
                line = line.strip()
                if not line: continue 
                if "\t" in line: line = line.replace("\t", ",") 
                if "," not in line: continue 
                    
                parts = line.split(',')
                r_name = parts[0].strip()
                r_dist_str = parts[-1].lower().replace('m', '').strip()
                
                try:
                    r_dist = float(r_dist_str)
                except ValueError:
                    continue 
                
                config = get_room_configuration(r_dist)
                if config:
                    valid_rooms.append({'name': r_name, 'distance': r_dist, 'config': config})
            
            if len(valid_rooms) == 0:
                lbl_status.configure(text="Error: No valid rooms found.\nFormat: 'Name, Distance'", text_color="#FF5555")
                return

            filepath = generate_multi_room_proposal(client, valid_rooms)
            lbl_status.configure(text=f"SUCCESS!\nSaved to Desktop/Alder_Quotes:\n{os.path.basename(filepath)}", text_color="#009A44")
            
            try: os.startfile(os.path.dirname(filepath))
            except: pass
            
        except PermissionError:
            lbl_status.configure(text="ERROR: PERMISSION DENIED.\nClose Word and try again.", text_color="#FF5555")
        except Exception as e:
            lbl_status.configure(text=f"Error: {str(e)}", text_color="#FF5555")

    # --- LAYOUT CONSTRUCTION ---
    
    # Main Container Frame (Card style - White)
    main_frame = ctk.CTkFrame(app, corner_radius=15, fg_color="white")
    main_frame.pack(pady=20, padx=20, fill="both", expand=True)

    # Logo Display
    if logo_image:
        lbl_logo = ctk.CTkLabel(main_frame, image=logo_image, text="")
        lbl_logo.pack(pady=(20, 5))
    else:
        # Fallback if image missing
        lbl_title = ctk.CTkLabel(main_frame, text="ALDER TECHNOLOGY", font=("Arial Black", 24), text_color="#009A44")
        lbl_title.pack(pady=(20, 5))

    lbl_subtitle = ctk.CTkLabel(main_frame, text="Multi-Room Proposal Generator", font=("Roboto", 13), text_color="gray")
    lbl_subtitle.pack(pady=(0, 20))

    # Input Section
    entry_client = ctk.CTkEntry(main_frame, placeholder_text="Client Name (e.g. Acme Corp)", width=350, height=40)
    entry_client.pack(pady=10)

    lbl_instr = ctk.CTkLabel(main_frame, text="Paste Excel Data below (Name | Depth)", text_color="gray", font=("Roboto", 12))
    lbl_instr.pack(pady=(10, 2))

    # Textbox with border
    txt_rooms = ctk.CTkTextbox(main_frame, width=350, height=200, corner_radius=10, border_width=1, border_color="#CCC")
    txt_rooms.pack(pady=5)
    
    # Generate Button
    btn_generate = ctk.CTkButton(
        main_frame, 
        text="GENERATE PROPOSAL", 
        command=on_generate_click, 
        width=350, 
        height=50, 
        font=("Roboto", 15, "bold"),
        fg_color="#009A44",  # Specific Alder Green
        hover_color="#007a36",
        text_color="white"
    )
    btn_generate.pack(pady=20)

    # Status Bar
    lbl_status = ctk.CTkLabel(main_frame, text="Ready", text_color="black", font=("Roboto", 12))
    lbl_status.pack(pady=10)
    
    # Footer
    lbl_footer = ctk.CTkLabel(app, text="v2.1 | Alder Technology Internal", text_color="#666", font=("Arial", 10))
    lbl_footer.pack(side="bottom", pady=10)

    app.mainloop()

# --- CRASH CATCHER ---
except Exception as e:
    root = tkinter.Tk()
    root.withdraw()
    error_msg = str(traceback.format_exc())
    tkinter.messagebox.showerror("Program Error", f"The program crashed:\n\n{error_msg}")
