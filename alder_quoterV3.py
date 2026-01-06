import customtkinter as ctk
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import sys

# ==========================================
# PART 1: THE LOGIC
# ==========================================

def get_room_configuration(distance):
    # --- TIER 1: UP TO 3.0M ---
    if distance <= 3.0:
        return {
            "tier_name": "Small Meeting Space",
            "cisco_core": "Cisco Room Bar",
            "cisco_cam": "Integrated Camera",
            "cisco_audio": "Integrated Audio",
            "display_model": "LG 55UL3J-B",
            "display_price": 1100.00,
            "mount_model": "Venturi VP-F80",
            "mount_price": 65.00,
            "service_price": 3000.00,
            "ms_annual": 1200.00
        }
    # --- TIER 2: UP TO 4.5M ---
    elif distance <= 4.5:
        return {
            "tier_name": "Medium Meeting Space",
            "cisco_core": "Cisco Room Bar",
            "cisco_cam": "Integrated Camera",
            "cisco_audio": "Table Mic Pro",
            "display_model": "LG 65UL3J-B",
            "display_price": 1500.00,
            "mount_model": "Venturi VP-F80",
            "mount_price": 65.00,
            "service_price": 3000.00,
            "ms_annual": 1200.00
        }
    # --- TIER 3: UP TO 5.5M ---
    elif distance <= 5.5:
        return {
            "tier_name": "Large Meeting Space",
            "cisco_core": "Cisco Room Bar Pro",
            "cisco_cam": "Dual Camera",
            "cisco_audio": "Ceiling Mic Pro",
            "display_model": "LG 75UL3J-B",
            "display_price": 2200.00,
            "mount_model": "Venturi VP-F80",
            "mount_price": 65.00,
            "service_price": 3000.00,
            "ms_annual": 1500.00
        }
    # --- TIER 4: UP TO 6.5M ---
    elif distance <= 6.5:
        return {
            "tier_name": "Extra Large Space",
            "cisco_core": "Cisco Room Bar Pro",
            "cisco_cam": "Dual Camera",
            "cisco_audio": "Ceiling Mic Pro",
            "display_model": "LG 86UL3J-B",
            "display_price": 3300.00,
            "mount_model": "VP-F100",
            "mount_price": 100.00,
            "service_price": 3500.00,
            "ms_annual": 1500.00
        }
    # --- TIER 5: UP TO 7.5M ---
    elif distance <= 7.5:
        return {
            "tier_name": "Boardroom",
            "cisco_core": "Cisco Kit EQ",
            "cisco_cam": "Quad Cam",
            "cisco_audio": "2x Mic + 6x Spk",
            "display_model": "LG 98UM5K",
            "display_price": 7500.00,
            "mount_model": "VP-F100",
            "mount_price": 100.00,
            "service_price": 4000.00,
            "ms_annual": 3000.00
        }
    else:
        return None

def generate_multi_room_proposal(client_name, room_list):
    doc = Document()
    
    # Document Styling
    section = doc.sections[0]
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # 1. HEADER
    header = doc.add_heading(f'AV Proposal: {client_name}', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Prepared by: Alder Technology')
    doc.add_paragraph(f'Date: {datetime.now().strftime("%d/%m/%Y")}')
    doc.add_paragraph(f'Total Rooms Scoped: {len(room_list)}')
    doc.add_paragraph('------------------------------------------------------')

    # 2. EXECUTIVE SUMMARY
    doc.add_heading('Partnership Overview', level=2)
    doc.add_paragraph(
        "Alder Technology is pleased to partner with Data#3 to provide this solution. "
        "The table below details the hardware and services required for each room provided in the scope.\n"
        "Please note: All Cisco hardware is listed for reference but is to be supplied and priced by Data#3."
    )

    # 3. MASTER SCHEDULE TABLE
    doc.add_heading('Master Room Schedule & Pricing', level=2)
    
    # Create Table with 6 Columns
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # Headers
    hdr = table.rows[0].cells
    hdr[0].text = "Room Name"
    hdr[1].text = "Cisco Type (Data3 Supply)"
    hdr[2].text = "Alder Display"
    hdr[3].text = "Alder Hardware"
    hdr[4].text = "Services + MS (5Yr)"
    hdr[5].text = "Room Total (Ex GST)"

    total_project_value = 0

    # LOOP THROUGH EVERY ROOM
    for room in room_list:
        name = room['name']
        dist = room['distance']
        data = room['config']
        
        # Calculate Costs
        hw_cost = data['display_price'] + data['mount_price']
        svc_cost = data['service_price']
        ms_5yr = data['ms_annual'] * 5
        room_total = hw_cost + svc_cost + ms_5yr
        
        # Add to Grand Total
        total_project_value += room_total

        # Add Row to Table
        row = table.add_row().cells
        row[0].text = f"{name}\n({dist}m)"
        row[1].text = f"{data['cisco_core']}\n{data['cisco_audio']}"
        row[2].text = data['display_model']
        row[3].text = f"${hw_cost:,.0f}"
        row[4].text = f"${(svc_cost + ms_5yr):,.0f}"
        row[5].text = f"${room_total:,.2f}"

    # 4. GRAND TOTALS
    doc.add_paragraph("\n")
    p_total = doc.add_paragraph()
    p_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    runner = p_total.add_run(f"GRAND TOTAL PROJECT VALUE (EX GST): ${total_project_value:,.2f}")
    runner.bold = True
    runner.font.size = Pt(16)
    runner.font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph("(Includes Alder Hardware, Professional Services, and 5-Year Managed Services)")

    # --- SAVE TO DESKTOP LOGIC ---
    # This finds the Desktop path for Windows
    desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop') 
    
    # Create a subfolder "Alder_Quotes" if it doesn't exist
    save_folder = os.path.join(desktop, "Alder_Quotes")
    if not os.path.exists(save_folder):
        os.makedirs(save_folder)

    # Sanitize Filename
    safe_client_name = "".join([c for c in client_name if c.isalpha() or c.isdigit() or c==' ']).rstrip()
    timestamp = datetime.now().strftime("%H-%M-%S")
    filename = f"Alder_Quote_{safe_client_name}_{timestamp}.docx"
    
    full_path = os.path.join(save_folder, filename)
    
    doc.save(full_path)
    return full_path

# ==========================================
# PART 2: THE USER INTERFACE (GUI)
# ==========================================

ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("blue")

app = ctk.CTk()
app.geometry("600x700")
app.title("Alder Technology - Multi-Room Quoter")

def on_generate_click():
    lbl_status.configure(text="Processing...", text_color="white")
    
    # 1. Get Client Name
    client = entry_client.get()
    if not client:
        lbl_status.configure(text="Error: Enter Client Name", text_color="#FF5555")
        return

    # 2. Get Raw Text
    raw_text = txt_rooms.get("0.0", "end")
    lines = raw_text.split('\n')
    
    valid_rooms = []
    
    # 3. Parse the Text List
    try:
        for line in lines:
            line = line.strip()
            if not line: continue 
            
            # SUPPORT EXCEL PASTING (Convert Tabs to Commas)
            if "\t" in line:
                line = line.replace("\t", ",")
            
            # Ensure we have a comma
            if "," not in line:
                continue 
                
            parts = line.split(',')
            r_name = parts[0].strip()
            r_dist_str = parts[-1].lower().replace('m', '').strip()
            
            try:
                r_dist = float(r_dist_str)
            except ValueError:
                continue # Skip lines where distance isn't a number
            
            # Get Config
            config = get_room_configuration(r_dist)
            
            if config:
                valid_rooms.append({
                    'name': r_name,
                    'distance': r_dist,
                    'config': config
                })
        
        if len(valid_rooms) == 0:
            lbl_status.configure(text="Error: No valid rooms found.\nFormat: 'Name, Distance'", text_color="#FF5555")
            return

        # 4. Generate Document
        filepath = generate_multi_room_proposal(client, valid_rooms)
        
        # Show Success on Screen
        lbl_status.configure(text=f"SUCCESS!\nFile saved to your Desktop:\n{filepath}", text_color="#00FF00")
        
        # Try to open the folder automatically
        os.startfile(os.path.dirname(filepath))
        
    except PermissionError:
        lbl_status.configure(text="ERROR: PERMISSION DENIED.\nClose Word and try again.", text_color="#FF5555")
    except Exception as e:
        lbl_status.configure(text=f"Error: {str(e)}", text_color="#FF5555")

# --- UI ELEMENTS ---

lbl_title = ctk.CTkLabel(app, text="Alder Multi-Room Quoter", font=("Roboto Medium", 24))
lbl_title.pack(pady=15)

entry_client = ctk.CTkEntry(app, placeholder_text="Client Name (e.g. Acme Corp)", width=400)
entry_client.pack(pady=5)

lbl_instr = ctk.CTkLabel(app, text="Paste Excel Data below (Name | Depth)", text_color="gray")
lbl_instr.pack(pady=(15, 5))

txt_rooms = ctk.CTkTextbox(app, width=400, height=300)
txt_rooms.pack(pady=5)

btn_generate = ctk.CTkButton(app, text="GENERATE MASTER PROPOSAL", command=on_generate_click, width=400, height=50, font=("Roboto Medium", 14))
btn_generate.pack(pady=20)

lbl_status = ctk.CTkLabel(app, text="Ready...", font=("Roboto", 12))
lbl_status.pack(pady=10)

app.mainloop()
