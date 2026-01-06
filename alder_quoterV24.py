import traceback
import tkinter
import tkinter.messagebox
import sys
import os
from datetime import datetime

# --- CRASH REPORTER WRAPPER ---
try:
    import customtkinter as ctk
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from PIL import Image
    import openpyxl 

    # ==========================================
    # PART 1: LOGIC & EXCEL DATABASE
    # ==========================================

    def load_pricelist_from_excel():
        """ Reads master_pricelist.xlsx and returns a list of dictionaries. """
        script_dir = os.path.dirname(os.path.abspath(__file__))
        file_path = os.path.join(script_dir, "master_pricelist.xlsx")
        
        if not os.path.exists(file_path):
            return None, "File 'master_pricelist.xlsx' not found."

        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            sheet = wb.active
            pricelist = []

            # Skip header (row 1), start from row 2
            for row in sheet.iter_rows(min_row=2, values_only=True):
                if row[0] is None: continue

                # Parse Cisco Items
                cisco_raw = row[2] if row[2] else ""
                cisco_list = [x.strip() for x in str(cisco_raw).split(',')]

                entry = {
                    "max_distance": float(row[0]),
                    "tier_name": str(row[1]),
                    "cisco_items": cisco_list,
                    "display_model": str(row[3]),
                    "display_price": float(row[4] or 0),
                    "mount_model": str(row[5]),
                    "mount_price": float(row[6] or 0),
                    "cables_misc": str(row[7]),
                    "cables_price": float(row[8] or 0),
                    "service_price": float(row[9] or 0),
                    "ms_annual": float(row[10] or 0)
                }
                pricelist.append(entry)

            pricelist.sort(key=lambda x: x['max_distance'])
            return pricelist, None

        except Exception as e:
            return None, str(e)

    def get_room_configuration(distance, pricelist):
        if not pricelist: return None
        for tier in pricelist:
            if distance <= tier['max_distance']:
                return tier
        return None

    def generate_multi_room_proposal(client_name, room_list, project_mode):
        
        # --- 1. SELECT TEMPLATE BASED ON MODE ---
        script_dir = os.path.dirname(os.path.abspath(__file__))
        
        if project_mode == "Data#3 (Cisco)":
            template_filename = "Template_Data3.docx"
            partner_name = "Data#3"
        else:
            template_filename = "Template_Fitout.docx"
            partner_name = "Client"

        template_path = os.path.join(script_dir, template_filename)
        
        if os.path.exists(template_path):
            doc = Document(template_path)
            doc.add_page_break() 
        else:
            # Fallback if specific template missing
            doc = Document()
            print(f"Warning: {template_filename} not found. Using blank.")

        # --- SET FONT TO HELVETICA ---
        try:
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Helvetica'
            font.size = Pt(10)
        except: pass 

        try:
            section = doc.sections[0]
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
        except: pass

        # --- HELPERS ---
        def shade_cell(cell, color_hex):
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
            cell._tc.get_or_add_tcPr().append(shading_elm)

        def add_manual_heading(text, size, color_rgb=None):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Helvetica'
            run.font.size = Pt(size)
            if color_rgb: run.font.color.rgb = color_rgb
            return p
            
        def format_row(row, height_cm):
            row.height = Cm(height_cm)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # ---------------------------------------------------------
        # PAGE 1: EXECUTIVE SUMMARY
        # ---------------------------------------------------------
        
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_t = p_title.add_run(f'AV Proposal: {client_name}')
        run_t.bold = True
        run_t.font.size = Pt(26)
        
        doc.add_paragraph(f'Prepared by: Alder Technology')
        doc.add_paragraph(f'Date: {datetime.now().strftime("%d/%m/%Y")}')
        doc.add_paragraph(f'Project Type: {project_mode}')
        doc.add_paragraph('------------------------------------------------------')

        add_manual_heading('Partnership Overview', 14)
        
        if project_mode == "Data#3 (Cisco)":
            overview_text = (
                "Alder Technology is pleased to partner with Data#3 to provide this solution. "
                "This document is split into two sections:\n"
                "1. A Master Financial Summary (Hardware + Year 1 Services).\n"
                "2. Detailed Bill of Materials for each specific room.\n\n"
                "Please note: Cisco hardware is listed for engineering reference but is to be supplied and priced by Data#3."
            )
        else:
            overview_text = (
                "Alder Technology is pleased to provide this comprehensive Audio Visual proposal. "
                "This document outlines the Master Financial Summary and the Detailed Bill of Materials for every room."
            )
            
        doc.add_paragraph(overview_text)

        add_manual_heading('1. Master Room Summary & Pricing', 14)
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        hdr_row = table.rows[0]
        format_row(hdr_row, 1.0)
        hdr = hdr_row.cells
        
        hdr[0].text = "Room Name"
        hdr[1].text = "Classification"
        hdr[2].text = "Supply & Services" 
        hdr[3].text = "Managed Service (Year 1)"
        hdr[4].text = "Total Year 1 (Ex GST)"

        for cell in hdr:
            shade_cell(cell, "D9E2F3")
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        grand_total_project = 0

        for room in room_list:
            name = room['name']
            dist = room['distance']
            data = room['config']
            
            # --- CALCULATE TOTALS BASED ON MODE ---
            upfront_cost = data['display_price'] + data['mount_price'] + data['cables_price'] + data['service_price']
            
            # If Fit-Out, we need to add the cost of the conferencing gear (which is currently not in Excel)
            # NOTE: Your Excel currently only has columns for Alder costs. 
            # Ideally, you would add a column in Excel for "VC Hardware Price". 
            # For now, I will assume 0 for VC hardware in Fit-Out mode unless you add it to Excel.
            
            ms_annual = data['ms_annual'] 
            
            room_total_y1 = upfront_cost + ms_annual
            grand_total_project += room_total_y1

            row_obj = table.add_row()
            format_row(row_obj, 0.9)
            row = row_obj.cells
            
            row[0].text = name
            row[1].text = f"{data['tier_name']} ({dist}m)"
            row[2].text = f"${upfront_cost:,.0f}"
            row[3].text = f"${ms_annual:,.0f}"
            row[4].text = f"${room_total_y1:,.2f}"
            
            for i in range(2, 5):
                row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph("\n")
        p_total = doc.add_paragraph()
        p_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        runner = p_total.add_run(f"TOTAL YEAR 1 PROJECT VALUE (EX GST): ${grand_total_project:,.2f}")
        runner.bold = True
        runner.font.size = Pt(16)
        runner.font.color.rgb = RGBColor(0, 102, 204)
        
        # ---------------------------------------------------------
        # PAGE 2+: DETAILED ROOM BREAKDOWNS
        # ---------------------------------------------------------
        
        doc.add_page_break()
        add_manual_heading('2. Detailed Room Specifications', 16)

        for room in room_list:
            name = room['name']
            data = room['config']
            
            # --- PROCESS ITEMS ---
            final_cisco_list = []
            moved_to_alder_list = []

            for item in data['cisco_items']:
                if not item: continue
                clean_item = item.strip()

                # If Data#3 Mode: Shure goes to Alder. 
                # If Fit-Out Mode: EVERYTHING is Alder.
                if project_mode == "Fit-Out (Full Scope)":
                    moved_to_alder_list.append(clean_item) # Everything goes to main list
                else:
                    # DATA#3 MODE LOGIC
                    if "Shure" in clean_item:
                        moved_to_alder_list.append(clean_item)
                    else:
                        # It is Cisco/Partner Scope
                        item_desc = clean_item
                        if "Room Bar Pro" in clean_item: item_desc += " / CS-BARPRO-K9"
                        elif "Room Bar" in clean_item and "Pro" not in clean_item: item_desc += " / CS-BAR-T-C-K9"
                        elif "Ceiling Microphone Pro" in clean_item: item_desc += " / CS-MIC-CLGPRO="
                        elif "Wire Hanging" in clean_item: item_desc += " / CS-MIC-CLGP-WHK="
                        final_cisco_list.append(item_desc)

            # --- TABLE SETUP ---
            table_room = doc.add_table(rows=0, cols=5)
            table_room.style = 'Table Grid'
            table_room.autofit = False 
            table_room.columns[0].width = Cm(3.0)
            table_room.columns[1].width = Cm(1.3)
            table_room.columns[2].width = Cm(8.5)
            table_room.columns[3].width = Cm(2.5)
            table_room.columns[4].width = Cm(2.7)
            
            # Header
            row_hdr_obj = table_room.add_row()
            format_row(row_hdr_obj, 1.0)
            row_hdr = row_hdr_obj.cells
            row_hdr[0].merge(row_hdr[4])
            row_hdr[0].text = f"ROOM: {name} - {data['tier_name']}"
            shade_cell(row_hdr[0], "1F4E79")
            row_hdr[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            row_hdr[0].paragraphs[0].runs[0].bold = True

            # Image Placeholder
            row_img_obj = table_room.add_row()
            format_row(row_img_obj, 4.0)
            row_img = row_img_obj.cells
            row_img[0].merge(row_img[4])
            row_img[0].text = "[PASTE FLOOR PLAN IMAGE HERE]"
            row_img[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Col Headers
            row_cols_obj = table_room.add_row()
            format_row(row_cols_obj, 0.9)
            row_cols = row_cols_obj.cells
            row_cols[0].text = "Item"
            row_cols[1].text = "Qty"
            row_cols[2].text = "Description / Model"
            row_cols[3].text = "Unit Cost"
            row_cols[4].text = "Total"
            for c in row_cols:
                shade_cell(c, "D9E2F3")
                c.paragraphs[0].runs[0].bold = True

            # --- SECTION 1: PARTNER SCOPE (ONLY FOR DATA#3 MODE) ---
            if project_mode == "Data#3 (Cisco)" and final_cisco_list:
                row_sec1_obj = table_room.add_row()
                format_row(row_sec1_obj, 0.8)
                row_sec1 = row_sec1_obj.cells
                row_sec1[0].merge(row_sec1[4])
                row_sec1[0].text = "1. Data#3 Supply Scope (Cisco Hardware)"
                shade_cell(row_sec1[0], "E7E6E6")
                row_sec1[0].paragraphs[0].runs[0].bold = True

                for item in final_cisco_list:
                    r_obj = table_room.add_row()
                    format_row(r_obj, 0.9)
                    r = r_obj.cells
                    for cell in r: shade_cell(cell, "FFF2CC")
                    r[0].text = "Video Conf"
                    r[1].text = "1"
                    r[2].text = item 
                    r[3].text = "Excl."
                    r[4].text = "Excl."
                    r[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    r[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # --- SECTION 2: ALDER SCOPE ---
            row_sec2_obj = table_room.add_row()
            format_row(row_sec2_obj, 0.8)
            row_sec2 = row_sec2_obj.cells
            row_sec2[0].merge(row_sec2[4])
            
            # Rename section based on mode
            if project_mode == "Fit-Out (Full Scope)":
                row_sec2[0].text = "1. Hardware & Services Scope"
            else:
                row_sec2[0].text = "2. Alder Technology Supply Scope"
                
            shade_cell(row_sec2[0], "E7E6E6")
            row_sec2[0].paragraphs[0].runs[0].bold = True

            def add_spec_row(item_cat, qty, desc_text, price_val):
                r_obj = table_room.add_row()
                format_row(r_obj, 0.9)
                r = r_obj.cells
                r[0].text = item_cat
                r[1].text = str(qty)
                r[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                r[2].text = desc_text
                r[3].text = f"${price_val:,.2f}"
                r[4].text = f"${(price_val * qty):,.2f}"
                r[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Add Standard Items
            add_spec_row("Visual Display", 1, data['display_model'], data['display_price'])
            add_spec_row("Mounting", 1, data['mount_model'], data['mount_price'])
            
            # Add Items from the moved list (Shure or Full Scope)
            for moved_item in moved_to_alder_list:
                # NOTE: If Fit-Out mode, these items currently have $0 cost because they come from the 'Cisco Items' column in Excel.
                # You may need to manually update Excel to include prices for these if doing full fit-outs.
                add_spec_row("Conf/Audio", 1, moved_item, 0.00) 

            add_spec_row("Cabling", 1, data['cables_misc'], data['cables_price'])
            
            # Service Row
            r_svc_obj = table_room.add_row()
            format_row(r_svc_obj, 0.9)
            r_svc = r_svc_obj.cells
            r_svc[0].text = "Services"
            r_svc[1].text = "1"
            r_svc[2].text = "Professional Services: Installation, Staging & PM"
            r_svc[3].text = f"${data['service_price']:,.2f}"
            r_svc[4].text = f"${data['service_price']:,.2f}"
            r_svc[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_svc[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # --- MANAGED SERVICES ---
            row_sec3_obj = table_room.add_row()
            format_row(row_sec3_obj, 0.8)
            row_sec3 = row_sec3_obj.cells
            row_sec3[0].merge(row_sec3[4])
            
            # Numbering changes based on mode
            if project_mode == "Fit-Out (Full Scope)":
                row_sec3[0].text = "2. Managed Services"
            else:
                row_sec3[0].text = "3. Managed Services"
                
            shade_cell(row_sec3[0], "E7E6E6")
            row_sec3[0].paragraphs[0].runs[0].bold = True

            r_msa_obj = table_room.add_row()
            format_row(r_msa_obj, 0.9)
            r_msa = r_msa_obj.cells
            r_msa[0].text = "Support"
            r_msa[1].text = "1"
            r_msa[2].text = "Managed Service Agreement - Year 1 (Annual Billing)"
            r_msa[3].text = f"${data['ms_annual']:,.2f}"
            r_msa[4].text = f"${data['ms_annual']:,.2f}"
            r_msa[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_msa[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # --- TOTAL ---
            upfront_subtotal = data['display_price'] + data['mount_price'] + data['cables_price'] + data['service_price']
            total_room_cost = upfront_subtotal + data['ms_annual']
            
            row_sub_obj = table_room.add_row()
            format_row(row_sub_obj, 1.0)
            row_sub = row_sub_obj.cells
            row_sub[0].merge(row_sub[3])
            row_sub[0].text = "TOTAL (EX GST):"
            row_sub[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_sub[0].paragraphs[0].runs[0].bold = True
            row_sub[4].text = f"${total_room_cost:,.2f}"
            row_sub[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_sub[4].paragraphs[0].runs[0].bold = True

            doc.add_paragraph("")

        # ---------------------------------------------------------
        # FINAL SECTION
        # ---------------------------------------------------------
        doc.add_page_break()
        add_manual_heading('Managed Service Agreement', 14)
        doc.add_paragraph(
            "Pricing excludes GST and is charged annually with an increase each year of 4% or CPI whichever is the greater. "
        )

        add_manual_heading('Exclusions', 14)
        exclusions_text = "Standard exclusions apply (Power, Data, Joinery, etc)."
        doc.add_paragraph(exclusions_text)

        doc.add_paragraph("\nThank you for your consideration.")
        doc.add_paragraph("Regards,")
        sig = doc.add_paragraph("George Coles")
        sig.runs[0].bold = True

        desktop = os.path.expanduser("~/Desktop")
        save_folder = os.path.join(desktop, "Alder_Quotes")
        if not os.path.exists(save_folder):
            os.makedirs(save_folder)

        safe_client_name = "".join([c for c in client_name if c.isalpha() or c.isdigit() or c==' ']).rstrip()
        timestamp = datetime.now().strftime("%H-%M-%S")
        filename = f"Alder_Quote_{safe_client_name}_{project_mode[:4]}_{timestamp}.docx"
        full_path = os.path.join(save_folder, filename)
        
        doc.save(full_path)
        return full_path

    # ==========================================
    # PART 2: THE USER INTERFACE (GUI)
    # ==========================================

    ctk.set_appearance_mode("Light") 
    ctk.set_default_color_theme("green") 

    app = ctk.CTk()
    app.geometry("550x900") 
    app.title("Alder Technology - Quoting Tool")
    app.configure(fg_color="#E5E5E5")

    PRICELIST_DATA, EXCEL_ERROR = load_pricelist_from_excel()
    ADDED_ROOMS = [] 

    script_dir = os.path.dirname(os.path.abspath(__file__))
    logo_path = os.path.join(script_dir, "alder_logo.png")
    
    logo_image = None
    if os.path.exists(logo_path):
        try:
            pil_image = Image.open(logo_path)
            logo_image = ctk.CTkImage(light_image=pil_image, dark_image=pil_image, size=(250, 180))
        except Exception: pass

    # --- REFRESH ROOM LIST ---
    def refresh_room_list():
        for widget in scroll_frame.winfo_children():
            widget.destroy()

        if not ADDED_ROOMS:
            lbl_empty = ctk.CTkLabel(scroll_frame, text="No rooms added yet.", text_color="gray")
            lbl_empty.pack(pady=10)
            return

        for index, room in enumerate(ADDED_ROOMS):
            row_frame = ctk.CTkFrame(scroll_frame, fg_color="transparent")
            row_frame.pack(fill="x", pady=2, padx=5)

            info_text = f"{room['name']} ({room['type']} - {room['distance']}m)"
            lbl_info = ctk.CTkLabel(row_frame, text=info_text, anchor="w", font=("Arial", 11))
            lbl_info.pack(side="left", padx=5)

            btn_del = ctk.CTkButton(
                row_frame, 
                text="X", 
                width=30, 
                height=20, 
                fg_color="#FF5555", 
                hover_color="darkred",
                command=lambda i=index: delete_room(i)
            )
            btn_del.pack(side="right", padx=5)

    def delete_room(index):
        if 0 <= index < len(ADDED_ROOMS):
            del ADDED_ROOMS[index]
            refresh_room_list()

    def on_add_room():
        r_name = entry_room_name.get().strip()
        r_dist_str = entry_distance.get().strip()
        r_type = dropdown_type.get()

        if not r_name:
            tkinter.messagebox.showwarning("Missing Data", "Please enter a Room Name.")
            return
        
        try:
            r_dist = float(r_dist_str)
        except ValueError:
            tkinter.messagebox.showwarning("Invalid Distance", "Distance must be a number.")
            return

        config = get_room_configuration(r_dist, PRICELIST_DATA)
        if not config:
            tkinter.messagebox.showerror("Error", f"Distance {r_dist}m exceeds Excel tiers.")
            return
        
        ADDED_ROOMS.append({'name': r_name, 'distance': r_dist, 'type': r_type, 'config': config})
        entry_room_name.delete(0, "end")
        refresh_room_list()

    def on_generate_click():
        lbl_status.configure(text="Processing...", text_color="black")
        app.update()
        
        if EXCEL_ERROR:
            lbl_status.configure(text=f"EXCEL ERROR: {EXCEL_ERROR}", text_color="#FF5555")
            return

        client = entry_client.get().strip()
        if not client:
            lbl_status.configure(text="Error: Enter Client Name", text_color="#FF5555")
            return

        if len(ADDED_ROOMS) == 0:
            lbl_status.configure(text="Error: Please add at least one room.", text_color="#FF5555")
            return

        # GET THE SELECTED MODE
        selected_mode = dropdown_project_mode.get()

        try:
            filepath = generate_multi_room_proposal(client, ADDED_ROOMS, selected_mode)
            lbl_status.configure(text=f"SUCCESS!\nSaved to Desktop/Alder_Quotes:\n{os.path.basename(filepath)}", text_color="#009A44")
            try: os.startfile(os.path.dirname(filepath))
            except: pass
        except PermissionError:
            lbl_status.configure(text="ERROR: PERMISSION DENIED.\nClose Word and try again.", text_color="#FF5555")
        except Exception as e:
            lbl_status.configure(text=f"Error: {str(e)}", text_color="#FF5555")

    # --- GUI LAYOUT ---
    
    main_frame = ctk.CTkFrame(app, corner_radius=15, fg_color="white")
    main_frame.pack(pady=20, padx=20, fill="both", expand=True)

    if logo_image:
        lbl_logo = ctk.CTkLabel(main_frame, image=logo_image, text="")
        lbl_logo.pack(pady=(20, 5))
    else:
        lbl_title = ctk.CTkLabel(main_frame, text="ALDER TECHNOLOGY", font=("Arial Black", 24), text_color="#009A44")
        lbl_title.pack(pady=(20, 5))

    # --- NEW: PROJECT MODE DROPDOWN ---
    lbl_mode = ctk.CTkLabel(main_frame, text="Select Project Mode:", font=("Roboto", 12, "bold"))
    lbl_mode.pack(pady=(5, 0))
    
    dropdown_project_mode = ctk.CTkOptionMenu(
        main_frame,
        values=["Data#3 (Cisco)", "Fit-Out (Full Scope)"],
        width=250,
        height=30,
        fg_color="#555555"
    )
    dropdown_project_mode.set("Data#3 (Cisco)")
    dropdown_project_mode.pack(pady=(0, 10))

    if EXCEL_ERROR:
        lbl_db_status = ctk.CTkLabel(main_frame, text=f"⚠️ {EXCEL_ERROR}", text_color="red", font=("Arial", 11, "bold"))
    else:
        lbl_db_status = ctk.CTkLabel(main_frame, text=f"✔ Loaded {len(PRICELIST_DATA)} tiers from Excel", text_color="green", font=("Arial", 11))
    lbl_db_status.pack(pady=(0, 10))

    entry_client = ctk.CTkEntry(main_frame, placeholder_text="Client Name (e.g. Acme Corp)", width=350, height=40)
    entry_client.pack(pady=5)

    ctk.CTkLabel(main_frame, text="------------------------------------------------", text_color="#CCC").pack()

    # --- ADD ROOM FORM ---
    tier_names = [item['tier_name'] for item in PRICELIST_DATA] if PRICELIST_DATA else ["Error loading Excel"]
    
    def on_dropdown_select(choice):
        if PRICELIST_DATA:
            for item in PRICELIST_DATA:
                if item['tier_name'] == choice:
                    entry_distance.delete(0, "end")
                    entry_distance.insert(0, str(item['max_distance']))
                    break

    dropdown_type = ctk.CTkOptionMenu(
        main_frame, 
        values=tier_names,
        width=350,
        height=35,
        command=on_dropdown_select,
        fg_color="#1F4E79",
        button_color="#143656"
    )
    dropdown_type.set("Select Room Type")
    dropdown_type.pack(pady=5)

    entry_room_name = ctk.CTkEntry(main_frame, placeholder_text="Room Name (e.g. Boardroom)", width=350, height=35)
    entry_room_name.pack(pady=5)

    frame_dist = ctk.CTkFrame(main_frame, fg_color="transparent")
    frame_dist.pack(pady=5)
    
    ctk.CTkLabel(frame_dist, text="Distance (m):", text_color="gray").pack(side="left", padx=5)
    entry_distance = ctk.CTkEntry(frame_dist, width=100, height=35)
    entry_distance.pack(side="left")

    btn_add = ctk.CTkButton(
        main_frame,
        text="+ Add Room to List",
        width=350,
        height=35,
        fg_color="#666666",
        hover_color="#444444",
        command=on_add_room
    )
    btn_add.pack(pady=10)

    lbl_list_title = ctk.CTkLabel(main_frame, text="Current Room List:", font=("Roboto", 12, "bold"))
    lbl_list_title.pack(pady=(10, 0))

    scroll_frame = ctk.CTkScrollableFrame(main_frame, width=350, height=150, corner_radius=10, border_width=1, border_color="#CCC")
    scroll_frame.pack(pady=5)
    
    refresh_room_list()

    btn_generate = ctk.CTkButton(
        main_frame, 
        text="GENERATE PROPOSAL", 
        command=on_generate_click, 
        width=350, 
        height=50, 
        font=("Roboto", 15, "bold"),
        fg_color="#009A44", 
        hover_color="#007a36",
        text_color="white",
        state="disabled" if EXCEL_ERROR else "normal"
    )
    btn_generate.pack(pady=20)

    lbl_status = ctk.CTkLabel(main_frame, text="Ready", text_color="black", font=("Roboto", 12))
    lbl_status.pack(pady=10)
    
    app.mainloop()

except Exception as e:
    root = tkinter.Tk()
    root.withdraw()
    error_msg = str(traceback.format_exc())
    tkinter.messagebox.showerror("Program Error", f"The program crashed:\n\n{error_msg}")
