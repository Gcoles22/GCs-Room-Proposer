import customtkinter as ctk
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

# ==========================================
# PART 1: THE LOGIC (Pricing & Tiers)
# ==========================================

def get_room_configuration(distance):
    # --- TIER 1: UP TO 3.0M ---
    if distance <= 3.0:
        return {
            "tier_name": "Small Meeting Space",
            "cisco_core": "Cisco Room Bar",
            "cisco_cam": "Integrated Camera",
            "cisco_audio": "Integrated Audio",
            "display_model": "LG 55UL3J-B",
            "display_price": 1100.00,
            "mount_model": "Venturi VP-F80",
            "mount_price": 65.00,
            "service_price": 3000.00,
            "ms_annual": 1200.00
        }
    # --- TIER 2: UP TO 4.5M ---
    elif distance <= 4.5:
        return {
            "tier_name": "Medium Meeting Space",
            "cisco_core": "Cisco Room Bar",
            "cisco_cam": "Integrated Camera",
            "cisco_audio": "Table Mic Pro",
            "display_model": "LG 65UL3J-B",
            "display_price": 1500.00,
            "mount_model": "Venturi VP-F80",
            "mount_price": 65.00,
            "service_price": 3000.00,
            "ms_annual": 1200.00
        }
    # --- TIER 3: UP TO 5.5M ---
    elif distance <= 5.5:
        return {
            "tier_name": "Large Meeting Space",
            "cisco_core": "Cisco Room Bar Pro",
            "cisco_cam": "Dual Camera",
            "cisco_audio": "Ceiling Mic Pro",
            "display_model": "LG 75UL3J-B",
            "display_price": 2200.00,
            "mount_model": "Venturi VP-F80",
            "mount_price": 65.00,
            "service_price": 3000.00,
            "ms_annual": 1500.00
        }
    # --- TIER 4: UP TO 6.5M ---
    elif distance <= 6.5:
        return {
            "tier_name": "Extra Large Space",
            "cisco_core": "Cisco Room Bar Pro",
            "cisco_cam": "Dual Camera",
            "cisco_audio": "Ceiling Mic Pro",
            "display_model": "LG 86UL3J-B",
            "display_price": 3300.00,
            "mount_model": "VP-F100",
            "mount_price": 100.00,
            "service_price": 3500.00,
            "ms_annual": 1500.00
        }
    # --- TIER 5: UP TO 7.5M ---
    elif distance <= 7.5:
        return {
            "tier_name": "Boardroom",
            "cisco_core": "Cisco Kit EQ",
            "cisco_cam": "Quad Cam",
            "cisco_audio": "2x Mic + 6x Spk",
            "display_model": "LG 98UM5K",
            "display_price": 7500.00,
            "mount_model": "VP-F100",
            "mount_price": 100.00,
            "service_price": 4000.00,
            "ms_annual": 3000.00
        }
    else:
        return None

def generate_multi_room_proposal(client_name, room_list):
    doc = Document()
    
    # --- DOCUMENT STYLING ---
    section = doc.sections[0]
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # 1. HEADER
    header = doc.add_heading(f'AV Proposal: {client_name}', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Prepared by: Alder Technology')
    doc.add_paragraph(f'Date: {datetime.now().strftime("%d/%m/%Y")}')
    doc.add_paragraph(f'Total Rooms Scoped: {len(room_list)}')
    doc.add_paragraph('------------------------------------------------------')

    # 2. EXECUTIVE SUMMARY
    doc.add_heading('Partnership Overview', level=2)
    doc.add_paragraph(
        "Alder Technology is pleased to partner with Data#3 to provide this solution. "
        "The table below details the hardware and services required for each room provided in the scope.\n"
        "Please note: All Cisco hardware is listed for reference but is to be supplied and priced by Data#3."
    )

    # 3. MASTER SCHEDULE TABLE
    doc.add_heading('Master Room Schedule & Pricing', level=2)
    
    # Create Table with 6 Columns
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # Headers
    hdr = table.rows[0].cells
    hdr[0].text = "Room Name"
    hdr[1].text = "Cisco Type (Data3 Supply)"
    hdr[2].text = "Alder Display"
    hdr[3].text = "Alder Hardware"
    hdr[4].text = "Services + MS (5Yr)"
    hdr[5].text = "Room Total (Ex GST)"

    total_project_value = 0

    # LOOP THROUGH EVERY ROOM
    for room in room_list:
        name = room['name']
        dist = room['distance']
        data = room['config']
        
        # Calculate Costs for this specific room
        hw_cost = data['display_price'] + data['mount_price']
        svc_cost = data['service_price']
        ms_5yr = data['ms_annual'] * 5
        room_total = hw_cost + svc_cost + ms_5yr
        
        # Add to Grand Total
        total_project_value += room_total

        # Add Row to Table
        row = table.add_row().cells
        row[0].text = f"{name}\n({dist}m)"
        row[1].text = f"{data['cisco_core']}\n{data['cisco_audio']}"
        row[2].text = data['display_model']
        row[3].text = f"${hw_cost:,.0f}"
        row[4].text = f"${(svc_cost + ms_5yr):,.0f}"
        row[5].text = f"${room_total:,.2f}"

    # 4. GRAND TOTALS
    doc.add_paragraph("\n")
    p_total = doc.add_paragraph()
    p_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    runner = p_total.add_run(f"GRAND TOTAL PROJECT VALUE (EX GST): ${total_project_value:,.2f}")
    runner.bold = True
    runner.font.size = Pt(16)
    runner.font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph("(Includes Alder Hardware, Professional Services, and 5-Year Managed Services for all rooms listed above)")

    # Save
    filename = f"Alder_Proposal_{client_name}_Master_Quote.docx"
    doc.save(filename)
    return filename

# ==========================================
# PART 2: THE USER INTERFACE (GUI)
# ==========================================

ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("blue")

app = ctk.CTk()
app.geometry("600x700")
app.title("Alder Technology - Multi-Room Quoter")

def on_generate_click():
    lbl_status.configure(text="Processing...", text_color="white")
    
    # 1. Get Client Name
    client = entry_client.get()
    if not client:
        lbl_status.configure(text="Error: Enter Client Name", text_color="#FF5555")
        return

    # 2. Get Raw Text from the Big Box
    raw_text = txt_rooms.get("0.0", "end")
    lines = raw_text.split('\n')
    
    valid_rooms = []
    
    # 3. Parse the Text List
    try:
        for line in lines:
            line = line.strip()
            if not line: continue # Skip empty lines
            
            if "," not in line:
                # Basic error handling for bad formatting
                continue 
                
            parts = line.split(',')
            r_name = parts[0].strip()
            # Clean the distance (remove 'm' if they typed 5m)
            r_dist_str = parts[1].lower().replace('m', '').strip()
            r_dist = float(r_dist_str)
            
            # Get Config
            config = get_room_configuration(r_dist)
            
            if config:
                valid_rooms.append({
                    'name': r_name,
                    'distance': r_dist,
                    'config': config
                })
        
        if len(valid_rooms) == 0:
            lbl_status.configure(text="Error: No valid rooms found.\nFormat: 'Room Name, Distance'", text_color="#FF5555")
            return

        # 4. Generate Document
        filename = generate_multi_room_proposal(client, valid_rooms)
        
        full_path = os.path.abspath(filename)
        lbl_status.configure(text=f"SUCCESS!\nProcessed {len(valid_rooms)} rooms.\nSaved: {filename}", text_color="#00FF00")
        
    except Exception as e:
        lbl_status.configure(text=f"Error parsing list: {str(e)}", text_color="#FF5555")

# --- UI ELEMENTS ---

lbl_title = ctk.CTkLabel(app, text="Alder Multi-Room Quoter", font=("Roboto Medium", 24))
lbl_title.pack(pady=15)

# Client Name
entry_client = ctk.CTkEntry(app, placeholder_text="Client Name (e.g. Acme Corp)", width=400)
entry_client.pack(pady=5)

# Instructions
lbl_instr = ctk.CTkLabel(app, text="Paste Room List below (Format: Name, Depth)\nExample:\nBoardroom, 7.5\nHuddle 1, 3.0", text_color="gray")
lbl_instr.pack(pady=(15, 5))

# Big Text Area for Rooms
txt_rooms = ctk.CTkTextbox(app, width=400, height=300)
txt_rooms.pack(pady=5)

# Generate Button
btn_generate = ctk.CTkButton(app, text="GENERATE MASTER PROPOSAL", command=on_generate_click, width=400, height=50, font=("Roboto Medium", 14))
btn_generate.pack(pady=20)

# Status Label
lbl_status = ctk.CTkLabel(app, text="Ready...", font=("Roboto", 12))
lbl_status.pack(pady=10)

app.mainloop()
